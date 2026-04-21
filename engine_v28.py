"""
engine.py — Banner Formatter core logic v2.8
All parsing, detection, and output writing lives here.
The Streamlit app (app.py) calls these functions directly.
"""

print("ENGINE VERSION 2.8 LOADED")

import io
import math
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import openpyxl
from openpyxl.styles import Font as XLFont, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter


# ── Helper functions ─────────────────────────────────────────

def lowerList(inputList):
    return [x.lower() if isinstance(x, str) else x for x in inputList]

def normal_round(n):
    if n - math.floor(n) < 0.5:
        return math.floor(n)
    return math.ceil(n)

def prettyPrint(number):
    if isinstance(number, str):
        return number
    try:
        if np.isnan(number):
            return ""
    except Exception:
        pass
    if isinstance(number, (int, float)):
        newNumber = str(normal_round(number))
        if len(newNumber) > 3:
            newNumber = newNumber[:-3] + ',' + newNumber[-3:]
        return newNumber
    return number

def removeNaN(theList):
    return [x for x in theList if not (isinstance(x, float) and math.isnan(x))]

def transposeList(theList):
    return pd.DataFrame(theList).T.values.tolist()

def sum2(listSum):
    total = 0
    for i in listSum:
        try:
            total += float(i)
        except Exception:
            pass
    return total


# ── Format detection ─────────────────────────────────────────

def detect_format(sheet_df):
    """
    Returns 'fmt2', 'fmt3', 'fmt4', 'fmt5', 'fmt6', 'fmt7', or 'fmt1'.
    fmt2 = standard question (countries/groups as columns, Total at row3 col0 or col1)
    fmt3 = top box summary (companies as rows, floating base)
    fmt4 = grid table (brands as columns, scale options as rows)
    fmt5 = weighted banner, question at row3
    fmt6 = fmt5 variant with extra descriptor row at row2
    fmt7 = weighted banner with group header row between question and categories
    fmt1 = original transposed format (legacy)
    """
    raw = sheet_df.values.tolist()

    # fmt6: descriptor at row2, question at row3, Total at row4 col1
    try:
        row2_col0 = raw[2][0]
        row3_col0 = raw[3][0]
        row4_col1 = raw[4][1]
        row5_col1 = raw[5][1]
        if (
            isinstance(row3_col0, str) and len(row3_col0.strip()) > 5
            and isinstance(row4_col1, str) and 'total' in row4_col1.lower()
            and isinstance(row5_col1, str) and row5_col1.strip()
            and (not isinstance(row2_col0, str) or len(row2_col0.strip()) < 50)
        ):
            return 'fmt6'
    except (IndexError, TypeError):
        pass

    # fmt5: question at row3, Total at row4 col1, category at row5 col1
    try:
        row3_col0 = raw[3][0]
        row4_col1 = raw[4][1]
        row5_col1 = raw[5][1]
        if (
            isinstance(row3_col0, str) and len(row3_col0.strip()) > 5
            and isinstance(row4_col1, str) and 'total' in row4_col1.lower()
            and isinstance(row5_col1, str) and row5_col1.strip()
        ):
            return 'fmt5'
    except (IndexError, TypeError):
        pass

    # fmt4: brands as columns — row3 col1 blank, row4 col1 is a brand name not 'Total'
    try:
        row3_col1 = raw[3][1]
        row4_col1 = raw[4][1]
        if (
            isinstance(row4_col1, str) and row4_col1.strip()
            and 'total' not in str(row4_col1).lower()
            and (
                row3_col1 is None
                or not isinstance(row3_col1, str)
                or row3_col1.strip() == ''
                or (isinstance(row3_col1, float) and math.isnan(row3_col1))
            )
        ):
            return 'fmt4'
    except (IndexError, TypeError):
        pass

    # fmt7: question at row2, group header at row3, categories at row4 col1=Total,
    #       Unweighted Base at row7
    try:
        row2_col0 = raw[2][0]
        row4_col1 = raw[4][1]
        row7_col0 = raw[7][0] if len(raw) > 7 else None
        if (
            isinstance(row2_col0, str) and len(row2_col0.strip()) > 10
            and isinstance(row4_col1, str) and 'total' in row4_col1.lower()
            and isinstance(row7_col0, str) and row7_col0.strip().lower().startswith('unweighted')
        ):
            return 'fmt7'
    except (IndexError, TypeError):
        pass

    # fmt2 / fmt3: row3 col0 OR col1 = 'Total', row4 = sub-labels
    # Handles two layouts:
    #   Layout A (standard):  col0=None,    col1='Total', col2='Belgium'...
    #   Layout B (Mastercard): col0='Total', col1='Belgium', col2='France'...
    try:
        row2_col0 = raw[2][0]
        row3_col0 = raw[3][0] if len(raw[3]) > 0 else None
        row3_col1 = raw[3][1] if len(raw[3]) > 1 else None
        row3_has_total = (
            (isinstance(row3_col1, str) and 'total' in row3_col1.lower()) or
            (isinstance(row3_col0, str) and 'total' in row3_col0.lower())
        )
        if (
            isinstance(row2_col0, str) and len(row2_col0.strip()) > 10
            and row3_has_total
        ):
            # Check rows 7-11 for a base label
            base_row_col0 = None
            base_row_col1 = None
            for check_row in range(7, min(12, len(raw))):
                cell = raw[check_row][0] if raw[check_row] else None
                if isinstance(cell, str):
                    stripped = cell.strip().lower()
                    if stripped.startswith('base') or stripped.startswith('unweighted'):
                        base_row_col0 = cell
                        base_row_col1 = raw[check_row][1] if len(raw[check_row]) > 1 else None
                        break
            if base_row_col0 is not None:
                if base_row_col1 is None or (isinstance(base_row_col1, float) and math.isnan(base_row_col1)):
                    return 'fmt3'
                return 'fmt2'
    except (IndexError, TypeError):
        pass

    return 'fmt1'


# ── Sheet scanning ───────────────────────────────────────────

def scan_file(file_bytes):
    """
    Read an uploaded Excel file and return metadata for every sheet.
    Returns list of dicts: {index, name, fmt, question_wording, columns}
    """
    xl = pd.ExcelFile(io.BytesIO(file_bytes))
    results = []
    for i, sheet_name in enumerate(xl.sheet_names):
        if i == 0:
            try:
                df0  = xl.parse(0, header=None, na_values=[''])
                fmt0 = detect_format(df0)
                if fmt0 == 'fmt1':
                    continue  # likely TOC/index
            except Exception:
                continue
        try:
            df  = xl.parse(i, header=None, na_values=[''])
            fmt = detect_format(df)
            raw = df.values.tolist()

            if fmt in ('fmt5', 'fmt6'):
                question_wording = str(raw[3][0]).strip() if (
                    len(raw) > 3 and raw[3] and raw[3][0]
                ) else sheet_name
            else:
                question_wording = str(raw[2][0]).strip() if (
                    len(raw) > 2 and raw[2] and raw[2][0]
                ) else sheet_name

            # Collect available columns
            # For fmt2/fmt3: detect whether Total is at col0 or col1
            columns = []
            if fmt in ('fmt2', 'fmt3'):
                group_row    = raw[3] if len(raw) > 3 else []
                sublabel_row = raw[4] if len(raw) > 4 else []
                # Layout B: col0 = 'Total', col1 = first country
                col0_val  = group_row[0] if group_row else None
                start_col = 0 if (isinstance(col0_val, str) and 'total' in col0_val.lower()) else 1
                for j in range(start_col, len(group_row)):
                    g = group_row[j]
                    s = sublabel_row[j] if j < len(sublabel_row) else ''
                    if isinstance(g, str) and g.strip():
                        columns.append((g.strip(), s.strip() if isinstance(s, str) else ''))
            elif fmt in ('fmt5', 'fmt6'):
                cat_row = raw[5] if len(raw) > 5 else []
                for j in range(1, len(cat_row)):
                    g = cat_row[j]
                    if isinstance(g, str) and g.strip() and g.strip() != '\xa0':
                        columns.append((g.strip(), ''))
            elif fmt == 'fmt7':
                cat_row = raw[4] if len(raw) > 4 else []
                for j in range(1, len(cat_row)):
                    g = cat_row[j]
                    if isinstance(g, str) and g.strip() and g.strip() != '\xa0':
                        columns.append((g.strip(), ''))
            elif fmt == 'fmt4':
                brand_row = raw[4] if len(raw) > 4 else []
                for j in range(1, len(brand_row)):
                    v = brand_row[j]
                    if isinstance(v, str) and v.strip():
                        columns.append((v.strip(), ''))

            results.append({
                'index':            i,
                'sheet_name':       sheet_name,
                'fmt':              fmt,
                'question_wording': question_wording,
                'columns':          columns,
            })
        except Exception as e:
            results.append({
                'index':            i,
                'sheet_name':       sheet_name,
                'fmt':              'error',
                'question_wording': f'Error reading sheet: {e}',
                'columns':          [],
            })
    return results


def get_all_columns(sheet_metas):
    seen = set()
    cols = []
    for s in sheet_metas:
        if s['fmt'] in ('fmt2', 'fmt3', 'fmt5', 'fmt6', 'fmt7'):
            for col in s['columns']:
                if col[0] not in seen:
                    seen.add(col[0])
                    cols.append(col)
    return cols


def scan_rows_for_sheets(file_bytes, sheet_indices):
    xl   = pd.ExcelFile(io.BytesIO(file_bytes))
    seen = set()
    rows = []
    for i in sheet_indices:
        try:
            df  = xl.parse(i, header=None, na_values=[''])
            fmt = detect_format(df)
            raw = df.values.tolist()
            if fmt == 'fmt2':
                base_row_idx = None
                for ri, row in enumerate(raw):
                    if isinstance(row[0], str) and row[0].strip().lower().startswith('base'):
                        base_row_idx = ri
                        break
                start = (base_row_idx + 2) if base_row_idx is not None else 9
                j = start
                while j < len(raw):
                    label = raw[j][0]
                    if isinstance(label, str) and label.strip():
                        clean = label.strip()
                        if clean.lower() == 'sigma':
                            break
                        if clean not in seen:
                            seen.add(clean)
                            rows.append(clean)
                    j += 3
            elif fmt in ('fmt5', 'fmt6', 'fmt7'):
                j = 11
                while j < len(raw):
                    label = raw[j][0]
                    if isinstance(label, str) and label.strip():
                        clean = label.strip()
                        if clean.lower() == 'sigma':
                            break
                        if clean not in seen:
                            seen.add(clean)
                            rows.append(clean)
                    j += 3
            elif fmt == 'fmt3':
                end_markers = {'overlap formula used', 'sigma'}
                j = 8
                while j < len(raw):
                    label = raw[j][0]
                    if isinstance(label, str) and label.strip():
                        clean = label.strip()
                        if clean.lower() in end_markers:
                            break
                        if clean not in seen:
                            seen.add(clean)
                            rows.append(clean)
                    j += 3
            elif fmt == 'fmt4':
                j = 9
                while j < len(raw):
                    label = raw[j][0]
                    if isinstance(label, str) and label.strip():
                        clean = label.strip()
                        if clean.lower() == 'sigma':
                            break
                        if clean not in seen:
                            seen.add(clean)
                            rows.append(clean)
                    j += 3
        except Exception:
            pass
    return rows


def get_question_groups(sheet_metas):
    import re
    groups = {}
    order  = []
    for m in sheet_metas:
        if m['fmt'] == 'error':
            continue
        match = re.match(r'^([A-Za-z]+\d+)', m['question_wording'].strip())
        if match:
            prefix = match.group(1).upper()
        else:
            prefix = m['sheet_name'][:6]
        if prefix not in groups:
            groups[prefix] = {'prefix': prefix, 'sheets': [], 'fmt': m['fmt']}
            order.append(prefix)
        groups[prefix]['sheets'].append(m)
    return [groups[p] for p in order]


# ── Parsers ──────────────────────────────────────────────────

def _select_cols(all_cols, desired_groups, sublabel_filter='total'):
    if desired_groups is not None:
        desired_lower = {d.lower() for d in desired_groups}
        return [(j, g, s) for (j, g, s) in all_cols if g.lower() in desired_lower]
    selected = [(j, g, s) for (j, g, s) in all_cols if s.lower() == sublabel_filter]
    if all_cols and all_cols[0] not in selected:
        selected.insert(0, all_cols[0])
    return selected


def parse_fmt2_sheet(sheet_df, desired_groups=None, weighted_data=False, weighted_base=None):
    if weighted_base is None:
        weighted_base = weighted_data
    raw = sheet_df.values.tolist()
    question_wording = str(raw[2][0]).strip() if raw[2][0] else ''

    group_row    = raw[3]
    sublabel_row = raw[4]
    all_cols     = []

    # Detect layout: Layout B has Total at col0, countries start at col1
    # Layout A has col0=answer label area, col1=Total, col2=first country
    col0_val  = group_row[0] if group_row else None
    start_col = 0 if (isinstance(col0_val, str) and 'total' in col0_val.lower()) else 1

    for j in range(start_col, len(group_row)):
        g = group_row[j]
        s = sublabel_row[j] if j < len(sublabel_row) else ''
        if isinstance(g, str) and g.strip():
            all_cols.append((j, g.strip(), s.strip() if isinstance(s, str) else ''))

    selected_cols = _select_cols(all_cols, desired_groups)
    col_indices   = [j for (j, g, s) in selected_cols]
    col_labels    = [(g, s) for (j, g, s) in selected_cols]

    # Base row
    base_row_idx = None
    for i, row in enumerate(raw):
        cell = row[0] if row else None
        if isinstance(cell, str):
            stripped = cell.strip().lower()
            if (stripped.startswith('base') or
                stripped.startswith('unweighted') or
                stripped.startswith('weighted base')):
                base_row_idx = i
                break

    base_offset   = 2 if weighted_base else 0
    base_values   = []
    if base_row_idx is not None:
        base_data_row = raw[base_row_idx + base_offset] if base_row_idx + base_offset < len(raw) else []
        base_values   = [base_data_row[j] if j < len(base_data_row) else None for j in col_indices]

    base_rows_used = 4 if weighted_data else 2
    data_start = (base_row_idx + base_rows_used) if base_row_idx is not None else 9

    answers  = []
    data     = []
    sig_data = []

    letter_row  = raw[5] if len(raw) > 5 else []
    col_letters = [letter_row[j] if j < len(letter_row) else None for j in col_indices]

    i = data_start
    while i < len(raw):
        label = raw[i][0]
        if isinstance(label, str) and label.strip():
            label_clean = label.strip()
            if label_clean.lower() == 'sigma':
                break
            answers.append(label_clean)
            pct_row = raw[i + 1] if i + 1 < len(raw) else []
            sig_row = raw[i + 2] if i + 2 < len(raw) else []
            row_vals = [pct_row[j] if j < len(pct_row) else None for j in col_indices]
            sig_vals = [sig_row[j]  if j < len(sig_row)  else None for j in col_indices]
            data.append(row_vals)
            sig_data.append(sig_vals)
        i += 3

    return {
        'question_wording': question_wording,
        'columns':          col_labels,
        'base_values':      base_values,
        'answers':          answers,
        'data':             data,
        'sig_data':         sig_data,
        'col_letters':      col_letters,
    }


def parse_fmt3_sheet(sheet_df, desired_groups=None):
    raw = sheet_df.values.tolist()
    question_wording = str(raw[2][0]).strip() if raw[2][0] else ''

    group_row    = raw[3]
    sublabel_row = raw[4]
    all_cols     = []
    col0_val  = group_row[0] if group_row else None
    start_col = 0 if (isinstance(col0_val, str) and 'total' in col0_val.lower()) else 1
    for j in range(start_col, len(group_row)):
        g = group_row[j]
        s = sublabel_row[j] if j < len(sublabel_row) else ''
        if isinstance(g, str) and g.strip():
            all_cols.append((j, g.strip(), s.strip() if isinstance(s, str) else ''))

    selected_cols = _select_cols(all_cols, desired_groups)
    col_indices   = [j for (j, g, s) in selected_cols]
    col_labels    = [(g, s) for (j, g, s) in selected_cols]

    end_markers = {'overlap formula used', 'sigma'}

    def get_val(row, j):
        v = row[j] if j < len(row) else None
        if v is None: return None
        if isinstance(v, str) and v.strip() in ('-', '', '\xa0'): return None
        if isinstance(v, float) and math.isnan(v): return None
        return v

    def is_numeric(v):
        return isinstance(v, (int, float)) and not (isinstance(v, float) and math.isnan(v))

    def row_has_nums(row):
        return any(is_numeric(row[j]) for j in col_indices if j < len(row))

    weighted_fmt3 = False
    data_start    = 8

    for check in range(7, min(16, len(raw))):
        row   = raw[check]
        label = row[0] if row else None
        if not isinstance(label, str) or not label.strip():
            continue
        stripped = label.strip().lower()
        if stripped.startswith('base') or stripped.startswith('unweighted'):
            continue
        if stripped in end_markers:
            break
        prev = raw[check - 1] if check > 0 else []
        if row_has_nums(prev):
            weighted_fmt3 = True
            data_start    = check - 1
        else:
            weighted_fmt3 = False
            data_start    = check
        break

    companies     = []
    company_bases = []
    data          = []

    i = data_start
    while i < len(raw):
        if weighted_fmt3:
            counts_row = raw[i]
            label_row  = raw[i + 1] if i + 1 < len(raw) else []
            label      = label_row[0] if label_row else None
            if not isinstance(label, str) or not label.strip():
                i += 1; continue
            label_clean = label.strip()
            if label_clean.lower() in end_markers: break
            companies.append(label_clean)
            company_bases.append([get_val(counts_row, j) for j in col_indices])
            data.append([get_val(label_row, j) for j in col_indices])
            i += 3
        else:
            label_row = raw[i]
            label     = label_row[0] if label_row else None
            if not isinstance(label, str) or not label.strip():
                i += 1; continue
            label_clean = label.strip()
            if label_clean.lower() in end_markers: break
            pct_row = raw[i + 1] if i + 1 < len(raw) else []
            companies.append(label_clean)
            company_bases.append([get_val(label_row, j) for j in col_indices])
            data.append([get_val(pct_row, j) for j in col_indices])
            i += 3

    return {
        'question_wording': question_wording,
        'columns':          col_labels,
        'base_values':      None,
        'company_bases':    company_bases,
        'answers':          companies,
        'data':             data,
    }


def parse_fmt4_sheet(sheet_df):
    raw = sheet_df.values.tolist()
    question_wording = str(raw[2][0]).strip() if raw[2][0] else ''

    brand_row     = raw[4]
    brands        = []
    brand_indices = []
    for j in range(1, len(brand_row)):
        v = brand_row[j]
        if isinstance(v, str) and v.strip():
            brands.append(v.strip())
            brand_indices.append(j)

    base_row    = raw[7]
    base_values = [base_row[j] if j < len(base_row) else None for j in brand_indices]

    answers = []
    data    = []
    i = 9
    while i < len(raw):
        label = raw[i][0]
        if isinstance(label, str) and label.strip():
            label_clean = label.strip()
            if label_clean.lower() == 'sigma': break
            answers.append(label_clean)
            pct_row  = raw[i + 1] if i + 1 < len(raw) else []
            row_vals = []
            for j in brand_indices:
                v = pct_row[j] if j < len(pct_row) else None
                if v is None or (isinstance(v, float) and math.isnan(v)):
                    row_vals.append(None)
                elif isinstance(v, str) and v.strip() in ('-', '', '\xa0'):
                    row_vals.append(None)
                else:
                    row_vals.append(v)
            data.append(row_vals)
        i += 3

    return {
        'question_wording': question_wording,
        'brands':           brands,
        'base_values':      base_values,
        'answers':          answers,
        'data':             data,
    }


def parse_fmt5_sheet(sheet_df, desired_groups=None, weighted_data=False, weighted_base=None):
    if weighted_base is None:
        weighted_base = weighted_data
    raw = sheet_df.values.tolist()
    question_wording = str(raw[3][0]).strip() if (len(raw) > 3 and raw[3][0]) else ''

    cat_row  = raw[5] if len(raw) > 5 else []
    all_cols = []
    for j in range(1, len(cat_row)):
        g = cat_row[j]
        if isinstance(g, str) and g.strip() and g.strip() != '\xa0':
            all_cols.append((j, g.strip(), ''))

    if desired_groups is not None:
        desired_lower = {d.lower() for d in desired_groups}
        selected_cols = [(j, g, s) for (j, g, s) in all_cols if g.lower() in desired_lower]
    else:
        selected_cols = all_cols

    col_indices = [j for (j, g, s) in selected_cols]
    col_labels  = [(g, s) for (j, g, s) in selected_cols]

    base_row_idx = 10 if weighted_base else 8
    base_row     = raw[base_row_idx] if len(raw) > base_row_idx else []
    base_values  = [base_row[j] if j < len(base_row) else None for j in col_indices]

    answers  = []
    data     = []
    sig_data = []

    letter_row  = raw[6] if len(raw) > 6 else []
    col_letters = [letter_row[j] if j < len(letter_row) else None for j in col_indices]

    i = 12
    while i < len(raw):
        label = raw[i][0]
        if isinstance(label, str) and label.strip():
            label_clean = label.strip()
            if label_clean.lower() == 'sigma': break
            answers.append(label_clean)
            pct_row  = raw[i + 1] if i + 1 < len(raw) else []
            sig_row  = raw[i + 2] if i + 2 < len(raw) else []
            row_vals, sig_vals = [], []
            for j in col_indices:
                v = pct_row[j] if j < len(pct_row) else None
                if v is None or (isinstance(v, float) and math.isnan(v)): row_vals.append(None)
                elif isinstance(v, str) and v.strip() in ('-', '', '\xa0'): row_vals.append(None)
                else: row_vals.append(v)
                sv = sig_row[j] if j < len(sig_row) else None
                sig_vals.append(sv if isinstance(sv, str) else None)
            data.append(row_vals)
            sig_data.append(sig_vals)
        i += 3

    return {
        'question_wording': question_wording,
        'columns':          col_labels,
        'base_values':      base_values,
        'answers':          answers,
        'data':             data,
        'sig_data':         sig_data,
        'col_letters':      col_letters,
    }


def parse_fmt6_sheet(sheet_df, desired_groups=None, weighted_data=False, weighted_base=None):
    if weighted_base is None:
        weighted_base = weighted_data
    raw = sheet_df.values.tolist()
    question_wording = str(raw[3][0]).strip() if (len(raw) > 3 and raw[3][0]) else ''

    cat_row  = raw[5] if len(raw) > 5 else []
    all_cols = []
    for j in range(1, len(cat_row)):
        g = cat_row[j]
        if isinstance(g, str) and g.strip() and g.strip() != '\xa0':
            all_cols.append((j, g.strip(), ''))

    if desired_groups is not None:
        desired_lower = {d.lower() for d in desired_groups}
        selected_cols = [(j, g, s) for (j, g, s) in all_cols if g.lower() in desired_lower]
    else:
        selected_cols = all_cols

    col_indices = [j for (j, g, s) in selected_cols]
    col_labels  = [(g, s) for (j, g, s) in selected_cols]

    base_row_idx = 10 if weighted_base else 8
    base_row     = raw[base_row_idx] if len(raw) > base_row_idx else []
    base_values  = [base_row[j] if j < len(base_row) else None for j in col_indices]

    letter_row  = raw[6] if len(raw) > 6 else []
    col_letters = [letter_row[j] if j < len(letter_row) else None for j in col_indices]

    answers  = []
    data     = []
    sig_data = []
    i = 12
    while i < len(raw):
        label = raw[i][0]
        if isinstance(label, str) and label.strip():
            label_clean = label.strip()
            if label_clean.lower() == 'sigma': break
            answers.append(label_clean)
            pct_row = raw[i + 1] if i + 1 < len(raw) else []
            sig_row = raw[i + 2] if i + 2 < len(raw) else []
            row_vals, sig_vals = [], []
            for j in col_indices:
                v = pct_row[j] if j < len(pct_row) else None
                if v is None or (isinstance(v, float) and math.isnan(v)): row_vals.append(None)
                elif isinstance(v, str) and v.strip() in ('-', '', '\xa0'): row_vals.append(None)
                else: row_vals.append(v)
                sv = sig_row[j] if j < len(sig_row) else None
                sig_vals.append(sv if isinstance(sv, str) else None)
            data.append(row_vals)
            sig_data.append(sig_vals)
        i += 3

    return {
        'question_wording': question_wording,
        'columns':          col_labels,
        'base_values':      base_values,
        'answers':          answers,
        'data':             data,
        'sig_data':         sig_data,
        'col_letters':      col_letters,
    }


def parse_fmt7_sheet(sheet_df, desired_groups=None, weighted_data=False, weighted_base=None):
    if weighted_base is None:
        weighted_base = weighted_data
    raw = sheet_df.values.tolist()
    question_wording = str(raw[2][0]).strip() if (len(raw) > 2 and raw[2][0]) else ""

    cat_row  = raw[4] if len(raw) > 4 else []
    all_cols = []
    for j in range(1, len(cat_row)):
        g = cat_row[j]
        if isinstance(g, str) and g.strip() and g.strip() != "\xa0":
            all_cols.append((j, g.strip(), ""))

    if desired_groups is not None:
        desired_lower = {d.lower() for d in desired_groups}
        selected_cols = [(j, g, s) for (j, g, s) in all_cols if g.lower() in desired_lower]
    else:
        selected_cols = all_cols

    col_indices = [j for (j, g, s) in selected_cols]
    col_labels  = [(g, s) for (j, g, s) in selected_cols]

    base_offset   = 2 if weighted_base else 0
    base_row_idx  = 7
    base_data_row = raw[base_row_idx + base_offset] if base_row_idx + base_offset < len(raw) else []
    base_values   = [base_data_row[j] if j < len(base_data_row) else None for j in col_indices]

    letter_row  = raw[5] if len(raw) > 5 else []
    col_letters = [letter_row[j] if j < len(letter_row) else None for j in col_indices]

    base_rows_used = 4 if weighted_data else 2
    data_start     = base_row_idx + base_rows_used

    answers, data, sig_data = [], [], []
    i = data_start
    while i < len(raw):
        label = raw[i][0]
        if isinstance(label, str) and label.strip():
            clean = label.strip()
            if any(clean.lower().startswith(b) for b in ("base", "unweighted", "sigma")):
                i += 1; continue
            answers.append(clean)
            pct_row = raw[i + 1] if i + 1 < len(raw) else []
            sig_row = raw[i + 2] if i + 2 < len(raw) else []
            row_vals, sig_vals = [], []
            for j in col_indices:
                v = pct_row[j] if j < len(pct_row) else None
                if v is None or (isinstance(v, float) and math.isnan(v)): row_vals.append(None)
                elif isinstance(v, str) and v.strip() in ("-", "", "\xa0"): row_vals.append(None)
                else: row_vals.append(v)
                sv = sig_row[j] if j < len(sig_row) else None
                sig_vals.append(sv if isinstance(sv, str) else None)
            data.append(row_vals)
            sig_data.append(sig_vals)
        i += 3

    return {
        "question_wording": question_wording,
        "columns":          col_labels,
        "base_values":      base_values,
        "answers":          answers,
        "data":             data,
        "sig_data":         sig_data,
        "col_letters":      col_letters,
    }


import re as _re

_FMT6_TABLE_TYPES = ['Summary Grid', 'Summary - Mean', 'T2B - Summary',
                     'B2B - Summary', 'T3B - Summary', 'HIDDEN']

def classify_fmt6_sheet(wording):
    w  = wording.replace("'[", "[")
    table_type_map = [
        ('summary_grid',  ['summary grid', 'grid - summary', 'grid summary']),
        ('summary_mean',  ['summary - mean', "summary - mean'"]),
        ('t2b',           ['t2b - summary', "t2b - 'summary", "t2b - summary'"]),
        ('b2b',           ['b2b - summary', "b2b - 'summary", "b2b - summary'"]),
        ('t3b',           ['t3b - summary', "t3b - 'summary"]),
        ('hidden',        ['hidden']),
    ]
    wl = w.lower()
    for key, patterns in table_type_map:
        for pat in patterns:
            if pat in wl:
                return key, None
    matches = _re.findall(r'\[([^\[\]]+)\]', w)
    for m in matches:
        ml = m.strip().lower()
        is_table = any(pat in ml for patterns in [p for _, p in table_type_map] for pat in patterns)
        if not is_table and len(ml) < 60:
            return 'entity', m.strip()
    return 'standalone', None


def parse_fmt6_mean(sheet_df, desired_groups=None, use_weighted_base=False):
    raw = sheet_df.values.tolist()
    question_wording = str(raw[3][0]).strip() if (len(raw) > 3 and raw[3][0]) else ''
    cat_row  = raw[5] if len(raw) > 5 else []
    all_cols = []
    for j in range(1, len(cat_row)):
        g = cat_row[j]
        if isinstance(g, str) and g.strip() and g.strip() != '\xa0':
            all_cols.append((j, g.strip()))
    if desired_groups:
        desired_lower = {d.lower() for d in desired_groups}
        selected = [(j, g) for j, g in all_cols if g.lower() in desired_lower]
    else:
        selected = all_cols
    col_indices = [j for j, g in selected]
    col_labels  = [g for j, g in selected]
    base_row  = raw[8] if len(raw) > 8 else []
    base_vals = [base_row[j] if j < len(base_row) else None for j in col_indices]
    answers, data = [], []
    SKIP = {'sigma', '- column means:', '- column proportions:', 'table of contents'}
    i = 10
    while i < len(raw):
        label = raw[i][0]
        if isinstance(label, str) and label.strip():
            clean = label.strip()
            if clean.lower() in SKIP: break
            answers.append(clean)
            row_vals = []
            for j in col_indices:
                v = raw[i][j] if j < len(raw[i]) else None
                if v is None or (isinstance(v, float) and math.isnan(v)): row_vals.append(None)
                elif isinstance(v, str) and v.strip() in ('-', '', '\xa0'): row_vals.append(None)
                else: row_vals.append(v)
            data.append(row_vals)
        i += 2
    return {
        'question_wording': question_wording,
        'columns':          [(g, '') for g in col_labels],
        'col_labels':       col_labels,
        'base_values':      base_vals,
        'answers':          answers,
        'data':             data,
        'sig_data':         [[] for _ in answers],
        'col_letters':      [],
        'is_mean':          True,
    }


def build_fmt6_entity_merge(entity_sheets, desired_groups, weighted_data, use_weighted_base):
    TARGET = {'top 2 box', 'top 2 box (net)', 'bottom 2 box', 'bottom 2 box (net)'}
    col_labels, base_values = [], []
    ref_answers, merged_data = None, None
    for sheet_idx, entity_name, sheet_df in entity_sheets:
        parsed = parse_fmt6_sheet(sheet_df, desired_groups, weighted_data, use_weighted_base)
        if not parsed['answers']: continue
        col_list = [g for (g, s) in parsed['columns']]
        if ref_answers is None:
            ref_answers = parsed['answers']
            merged_data = [[] for _ in ref_answers]
        ans_lookup = {a.strip().lower(): i for i, a in enumerate(parsed['answers'])}
        for ci, col_label in enumerate(col_list):
            col_labels.append(f"{entity_name}\n{col_label}")
            base_values.append(parsed['base_values'][ci] if ci < len(parsed['base_values']) else None)
            for ai, ref_ans in enumerate(ref_answers):
                idx = ans_lookup.get(ref_ans.strip().lower())
                val = parsed['data'][idx][ci] if (idx is not None and idx < len(parsed['data']) and ci < len(parsed['data'][idx])) else None
                merged_data[ai].append(val)
    if not ref_answers: return None
    filtered_a, filtered_d = [], []
    for ai, ans in enumerate(ref_answers):
        if any(kw in ans.strip().lower() for kw in TARGET):
            filtered_a.append(ans.strip())
            filtered_d.append(merged_data[ai])
    return (col_labels, base_values, filtered_a, filtered_d) if filtered_a else None


def _hex_to_rgb(hex_color):
    h = hex_color.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def _interpolate_color(start_hex, end_hex, t):
    r1,g1,b1 = _hex_to_rgb(start_hex)
    r2,g2,b2 = _hex_to_rgb(end_hex)
    r = int(r1 + (r2-r1)*t); g = int(g1 + (g2-g1)*t); b = int(b1 + (b2-b1)*t)
    return f"{r:02X}{g:02X}{b:02X}"

HEATMAP_SCHEMES = {
    "Blue scale":          ("#EFF6FF", "#1D4ED8"),
    "Green scale":         ("#F0FDF4", "#166534"),
    "Red-Green diverging": ("#DC2626", "#16A34A"),
    "None":                None,
}

def get_heatmap_color(value, min_val, max_val, scheme_name, custom_start=None, custom_end=None):
    if scheme_name == "None" or min_val is None or max_val is None: return None
    t = 0.5 if max_val == min_val else max(0.0, min(1.0, (value-min_val)/(max_val-min_val)))
    if scheme_name == "Custom" and custom_start and custom_end:
        start, end = custom_start, custom_end
    else:
        pair = HEATMAP_SCHEMES.get(scheme_name)
        if not pair: return None
        start, end = pair
    return _interpolate_color(start, end, t)


def build_sig_flags(sig_data, col_letters, total_col_idx=0):
    if not sig_data or not col_letters:
        return [['' for _ in row] for row in sig_data]
    total_letter = col_letters[total_col_idx] if total_col_idx < len(col_letters) else 'A'
    flags = []
    for row_sig in sig_data:
        row_flags = []
        total_sig_str = str(row_sig[total_col_idx]) if total_col_idx < len(row_sig) else ''
        for ci, sig_str in enumerate(row_sig):
            if ci == total_col_idx:
                row_flags.append(''); continue
            sig_str     = str(sig_str) if sig_str else ''
            this_letter = col_letters[ci] if ci < len(col_letters) else ''
            if total_letter and total_letter.upper() in sig_str.upper():
                row_flags.append(' ▲')
            elif this_letter and this_letter.upper() in total_sig_str.upper():
                row_flags.append(' ▼')
            else:
                row_flags.append('')
        flags.append(row_flags)
    return flags


# ── Word output ──────────────────────────────────────────────

def _get_word_template(portrait_landscape=False):
    import os
    template_file = 'template_portrait.docx' if portrait_landscape else 'template_landscape.docx'
    if os.path.exists(template_file):
        return Document(template_file)
    return Document()


def write_table_to_doc(doc, question_wording, col_labels, base_values,
                       answers, data, multiple, is_first,
                       sig_flags=None, heatmap_scheme=None,
                       heatmap_custom_start=None, heatmap_custom_end=None):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _set_cell_shading(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

    if not is_first:
        doc.add_paragraph()
    q_para       = doc.add_paragraph()
    q_para.style = doc.styles['Normal']
    q_para.text  = question_wording

    n_cols = len(col_labels)
    table  = doc.add_table(rows=1, cols=n_cols + 1)
    table.style     = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    hdr[0].text = ''
    for i in range(n_cols):
        base_str = prettyPrint(base_values[i]) if i < len(base_values) else ''
        hdr[i+1].text = ''
        hdr[i+1].paragraphs[0].add_run(f"{col_labels[i]}\n(N={base_str})").bold = True
        hdr[i+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    all_vals = []
    if heatmap_scheme and heatmap_scheme != "None":
        for row in data:
            for v in row:
                if v is not None and not isinstance(v, str):
                    try: all_vals.append(float(v) * multiple)
                    except Exception: pass
    min_val = min(all_vals) if all_vals else None
    max_val = max(all_vals) if all_vals else None

    for r_idx, answer in enumerate(answers):
        if str(answer).strip().lower() == 'sigma': continue
        row_cells = table.add_row().cells
        row_cells[0].text = str(answer)
        row_vals  = data[r_idx]      if r_idx < len(data)      else []
        row_flags = sig_flags[r_idx] if sig_flags and r_idx < len(sig_flags) else []
        for c_idx in range(n_cols):
            val  = row_vals[c_idx]  if c_idx < len(row_vals)  else None
            flag = row_flags[c_idx] if c_idx < len(row_flags) else ''
            cell = row_cells[c_idx + 1]
            cell.text = ''
            if val is None or (isinstance(val, float) and math.isnan(val)):
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif isinstance(val, str):
                cell.paragraphs[0].add_run(val + flag)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                pct_val = normal_round(round(val * multiple, 3))
                cell.paragraphs[0].add_run(f"{pct_val}%{flag}")
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if heatmap_scheme and heatmap_scheme != "None" and min_val is not None:
                    hex_col = get_heatmap_color(float(val)*multiple, min_val, max_val,
                                                heatmap_scheme, heatmap_custom_start, heatmap_custom_end)
                    if hex_col: _set_cell_shading(cell, hex_col)

    for j in range(n_cols + 1):
        for cell in table.columns[j].cells:
            cell.width = Inches(3) if j == 0 else Inches(1.75)
            if j > 0:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# ── Excel output ─────────────────────────────────────────────

HDR_FILL  = PatternFill("solid", fgColor="1F4E79")
HDR_FONT  = XLFont(bold=True, color="FFFFFF", name="Arial", size=10)
BODY_FONT = XLFont(name="Arial", size=10)
CTR       = Alignment(horizontal="center", vertical="center", wrap_text=True)
LEFT      = Alignment(horizontal="left",   vertical="center", wrap_text=True)
THIN      = Side(style="thin", color="AAAAAA")
BORDER    = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


def _build_xl_toc(wb, toc_entries):
    """
    Build a Table of Contents sheet as the FIRST sheet in the workbook.
    toc_entries: list of (question_wording, sheet_name)
    Each entry links to cell A1 of the named sheet.
    Also adds a '← Contents' back-link in cell A1 of every data sheet.
    """
    toc = wb.create_sheet(title='Contents', index=0)
    toc.sheet_view.showGridLines = False
    toc.column_dimensions['A'].width = 8
    toc.column_dimensions['B'].width = 72
    toc.column_dimensions['C'].width = 22

    hdr_fill = PatternFill('solid', fgColor='1F4E79')
    hdr_font = XLFont(bold=True, color='FFFFFF', name='Arial', size=11)
    thin     = Side(style='thin', color='CCCCCC')
    brd      = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col, val in [(1,'#'), (2,'Question'), (3,'Sheet')]:
        c = toc.cell(row=1, column=col, value=val)
        c.font=hdr_font; c.fill=hdr_fill
        c.alignment=Alignment(horizontal='left', vertical='center')
        c.border=brd
    toc.row_dimensions[1].height = 22

    link_font  = XLFont(color='0563C1', underline='single', name='Arial', size=10)
    plain_font = XLFont(name='Arial', size=10)
    alt_fill   = PatternFill('solid', fgColor='F2F7FC')
    ethin      = Side(style='thin', color='E0E0E0')
    ebrd       = Border(left=ethin, right=ethin, top=ethin, bottom=ethin)

    seen = set()
    toc_row = 2
    for i, (wording, sheet_name) in enumerate(toc_entries):
        if sheet_name in seen:
            continue
        seen.add(sheet_name)
        fill = alt_fill if i % 2 == 0 else PatternFill()
        safe = sheet_name.replace("'", "''")

        num_cell = toc.cell(row=toc_row, column=1, value=i+1)
        num_cell.font=plain_font; num_cell.border=ebrd; num_cell.fill=fill
        num_cell.alignment=Alignment(horizontal='center', vertical='center')

        q_cell = toc.cell(row=toc_row, column=2, value=wording[:120])
        q_cell.hyperlink = f"#{safe}!A1"
        q_cell.font=link_font; q_cell.border=ebrd; q_cell.fill=fill
        q_cell.alignment=Alignment(horizontal='left', vertical='center', wrap_text=True)

        s_cell = toc.cell(row=toc_row, column=3, value=sheet_name)
        s_cell.font=plain_font; s_cell.border=ebrd; s_cell.fill=fill
        s_cell.alignment=Alignment(horizontal='left', vertical='center')

        toc.row_dimensions[toc_row].height = 18
        toc_row += 1

    # Add back-link row at top of every data sheet
    back_font = XLFont(color='0563C1', underline='single', name='Arial', size=9)
    for ws in wb.worksheets:
        if ws.title == 'Contents':
            continue
        ws.insert_rows(1)
        nav = ws.cell(row=1, column=1, value='← Contents')
        nav.hyperlink = '#Contents!A1'
        nav.font      = back_font
        nav.alignment = Alignment(horizontal='left', vertical='center')

    return toc


def _xl_write_table(ws, start_row, question_wording, col_headers,
                    base_values, answers, data, multiple=100,
                    show_base_in_header=True, sig_flags=None,
                    heatmap_scheme=None, heatmap_custom_start=None,
                    heatmap_custom_end=None):
    ws.cell(row=start_row, column=1, value=question_wording)
    ws.cell(row=start_row, column=1).font = XLFont(bold=True, name="Arial", size=10)
    ws.merge_cells(start_row=start_row, start_column=1,
                   end_row=start_row, end_column=len(col_headers) + 1)
    start_row += 1

    ws.cell(row=start_row, column=1, value='').fill   = HDR_FILL
    ws.cell(row=start_row, column=1).border = BORDER
    for ci, label in enumerate(col_headers):
        n_str = ''
        if show_base_in_header and ci < len(base_values) and base_values[ci] is not None:
            n_str = f"\n(N={prettyPrint(base_values[ci])})"
        cell           = ws.cell(row=start_row, column=ci + 2, value=label + n_str)
        cell.font      = HDR_FONT
        cell.fill      = HDR_FILL
        cell.alignment = CTR
        cell.border    = BORDER
    start_row += 1

    all_vals = []
    if heatmap_scheme and heatmap_scheme != "None":
        for row in data:
            for v in row:
                if v is not None and not isinstance(v, str):
                    try: all_vals.append(float(v) * multiple)
                    except Exception: pass
    min_val = min(all_vals) if all_vals else None
    max_val = max(all_vals) if all_vals else None

    for ri, answer in enumerate(answers):
        if str(answer).strip().lower() == 'sigma': continue
        row_vals  = data[ri]      if ri < len(data)      else []
        row_flags = sig_flags[ri] if sig_flags and ri < len(sig_flags) else []
        lc = ws.cell(row=start_row, column=1, value=str(answer))
        lc.font=BODY_FONT; lc.alignment=LEFT; lc.border=BORDER
        for ci in range(len(col_headers)):
            val  = row_vals[ci]  if ci < len(row_vals)  else None
            flag = row_flags[ci] if ci < len(row_flags) else ''
            cell = ws.cell(row=start_row, column=ci + 2)
            cell.font=BODY_FONT; cell.alignment=CTR; cell.border=BORDER
            if val is None or (isinstance(val, float) and math.isnan(val)):
                cell.value = ''
            elif isinstance(val, str):
                cell.value = val + (flag or '')
            else:
                pct_val    = normal_round(round(val * multiple, 3))
                cell.value = f"{pct_val}%{flag or ''}"
                if heatmap_scheme and heatmap_scheme != "None" and min_val is not None:
                    hex_col = get_heatmap_color(float(val)*multiple, min_val, max_val,
                                                heatmap_scheme, heatmap_custom_start, heatmap_custom_end)
                    if hex_col: cell.fill = PatternFill("solid", fgColor=hex_col)
        start_row += 1

    ws.column_dimensions['A'].width = max(ws.column_dimensions['A'].width, 35)
    for ci in range(len(col_headers)):
        col_letter = get_column_letter(ci + 2)
        ws.column_dimensions[col_letter].width = max(ws.column_dimensions[col_letter].width, 14)

    return start_row + 1


def _prep_fmt2_fmt5(parsed, row_filter, show_sig_flags):
    col_labels  = [g for (g, s) in parsed['columns']]
    all_answers = parsed["answers"]
    all_data    = parsed["data"]
    all_sig     = parsed.get("sig_data", [[] for _ in all_answers])

    if row_filter and row_filter != 'all':
        custom   = {r.strip().lower() for r in row_filter}
        filtered = [(a, d, s) for a, d, s in zip(all_answers, all_data, all_sig)
                    if a.strip().lower() in custom]
        answers  = [f[0] for f in filtered]
        data     = [f[1] for f in filtered]
        sig_data = [f[2] for f in filtered]
    else:
        answers  = all_answers
        data     = all_data
        sig_data = all_sig

    flags = None
    if show_sig_flags and sig_data:
        flags = build_sig_flags(sig_data, parsed.get('col_letters', []))

    return col_labels, answers, data, flags


def _generate_fmt6_outputs(
    xl, file_bytes, sheet_table_configs, desired_groups,
    output_format, excel_mode, portrait_landscape,
    weighted_data, use_weighted_base,
    out_merged, out_t2b, out_b2b, out_grid, out_mean, out_standalone,
    progress_callback,
):
    skipped, errors = [], []
    groups, order = {}, []
    selected_indices = set()
    for cfg in sheet_table_configs:
        si = cfg[0] if not isinstance(cfg[0], str) else xl.sheet_names.index(cfg[0])
        selected_indices.add(si)

    for i in range(len(xl.sheet_names)):
        if i not in selected_indices: continue
        try:
            df  = xl.parse(i, header=None, na_values=[""])
            raw = df.values.tolist()
            wording = ""
            for row_idx in [3, 2]:
                if len(raw) > row_idx and raw[row_idx] and raw[row_idx][0]:
                    val = str(raw[row_idx][0]).strip()
                    if len(val) > 5 and "sample" not in val.lower() and "weight" not in val.lower():
                        wording = val; break
            if not wording: continue
            sheet_type, entity = classify_fmt6_sheet(wording)
            if sheet_type == "hidden": continue
            m       = _re.match(r"^([A-Za-z0-9_]+)\.", wording.strip())
            prefix  = m.group(1) if m else wording[:20]
            if prefix not in groups:
                groups[prefix] = {
                    "wording_base": _re.sub(r"\s*\[.*?\]\s*-?\s*", " ", wording).strip(),
                    "entity": [], "t2b": [], "b2b": [], "t3b": [],
                    "summary_grid": [], "summary_mean": [], "standalone": [],
                }
                order.append(prefix)
            entry = (i, entity, df)
            if sheet_type == "entity":        groups[prefix]["entity"].append(entry)
            elif "t2b" in sheet_type:         groups[prefix]["t2b"].append(entry)
            elif "b2b" in sheet_type:         groups[prefix]["b2b"].append(entry)
            elif sheet_type == "summary_grid":groups[prefix]["summary_grid"].append(entry)
            elif "mean" in sheet_type:        groups[prefix]["summary_mean"].append(entry)
            else:                             groups[prefix]["standalone"].append(entry)
        except Exception as e:
            errors.append((xl.sheet_names[i], str(e)))

    word_doc = None
    if output_format in ("word", "both"):
        try: word_doc = _get_word_template(portrait_landscape)
        except Exception: word_doc = Document()
        word_doc.styles["Normal"].font.name = "Arial"
        word_doc.styles["Normal"].font.size = Pt(10)

    xl_wb, xl_sheet_rows, xl_sheet_ctr = None, {}, [0]
    toc_entries = []
    if output_format in ("excel", "both"):
        xl_wb = openpyxl.Workbook()
        xl_wb.remove(xl_wb.active)

    def _xl_get(label):
        name = label[:31]
        if excel_mode == "per_question":
            if name not in xl_wb.sheetnames:
                xl_wb.create_sheet(title=name)
                xl_sheet_rows[name] = 1
            return xl_wb[name], xl_sheet_rows.get(name, 1)
        else:
            xl_sheet_ctr[0] += 1
            sname = f"{xl_sheet_ctr[0]:03d}_{name}"[:31]
            return xl_wb.create_sheet(title=sname), 1

    def _xl_done(ws, label, nr):
        if excel_mode == "per_question":
            xl_sheet_rows[label[:31]] = nr

    first_word = [True]

    def _write(q_wording, col_labels, base_values, answers, data, xl_key, is_mean=False):
        multiple = 1 if is_mean else 100
        if word_doc is not None:
            if not first_word[0]: word_doc.add_paragraph()
            qp = word_doc.add_paragraph()
            qp.style = word_doc.styles["Normal"]
            qp.text  = q_wording
            n_cols = len(col_labels)
            tbl = word_doc.add_table(rows=1, cols=n_cols + 1)
            tbl.style     = "Table Grid"
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = tbl.rows[0].cells
            hdr[0].text = ""
            for ci in range(n_cols):
                base_str = prettyPrint(base_values[ci]) if ci < len(base_values) else ""
                hdr[ci+1].text = ""
                run = hdr[ci+1].paragraphs[0].add_run(f"{col_labels[ci]}\n(N={base_str})")
                run.bold = True
                hdr[ci+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for ri, ans in enumerate(answers):
                rc = tbl.add_row().cells
                rc[0].text = str(ans)
                row_vals = data[ri] if ri < len(data) else []
                for ci in range(n_cols):
                    val  = row_vals[ci] if ci < len(row_vals) else None
                    cell = rc[ci + 1]; cell.text = ""
                    if val is None or (isinstance(val, float) and math.isnan(val)): pass
                    elif isinstance(val, str):
                        cell.paragraphs[0].add_run(val)
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif is_mean:
                        cell.paragraphs[0].add_run(str(round(val, 2)))
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        cell.paragraphs[0].add_run(f"{normal_round(round(val*100,3))}%")
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for j in range(n_cols + 1):
                for cell in tbl.columns[j].cells:
                    cell.width = Inches(2.5) if j == 0 else Inches(1.4)
                    if j > 0: cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            first_word[0] = False

        if xl_wb is not None:
            ws, row_num = _xl_get(xl_key)
            nr = _xl_write_table(ws, row_num, q_wording, col_labels, base_values, answers, data, multiple)
            _xl_done(ws, xl_key, nr)
            toc_entries.append((q_wording[:100], ws.title))

    total_groups = max(len(order), 1)
    for gi, prefix in enumerate(order):
        if progress_callback: progress_callback(gi / total_groups, f"Processing {prefix}…")
        grp    = groups[prefix]
        xl_key = prefix[:31]

        if out_merged and grp["entity"]:
            result = build_fmt6_entity_merge(grp["entity"], desired_groups, weighted_data, use_weighted_base)
            if result:
                cls, bvs, ans, dat = result
                _write(grp["wording_base"] + " [Companies T2B/B2B]", cls, bvs, ans, dat, xl_key)

        type_cfg = [
            ("t2b",          out_t2b,  False, " [T2B Summary]"),
            ("b2b",          out_b2b,  False, " [B2B Summary]"),
            ("summary_grid", out_grid, False, " [Summary Grid]"),
            ("summary_mean", out_mean, True,  " [Mean]"),
        ]
        for key, enabled, use_mean, suffix in type_cfg:
            if not enabled: continue
            for si, entity, df in grp[key]:
                try:
                    if use_mean:  p = parse_fmt6_mean(df, desired_groups, use_weighted_base)
                    elif key == "summary_grid": p = parse_fmt6_sheet(df, None, weighted_data, use_weighted_base)
                    else: p = parse_fmt6_sheet(df, desired_groups, weighted_data, use_weighted_base)
                    if not p["answers"]: continue
                    cls = p.get("col_labels") or [g for g,s in p["columns"]]
                    _write(p["question_wording"] + suffix, cls, p["base_values"],
                           p["answers"], p["data"], xl_key, is_mean=p.get("is_mean", False))
                except Exception as e:
                    errors.append((xl.sheet_names[si], str(e)))

        if out_standalone:
            for si, entity, df in grp["standalone"]:
                try:
                    p = parse_fmt6_sheet(df, desired_groups, weighted_data, use_weighted_base)
                    if not p["answers"]: continue
                    cls = p.get("col_labels") or [g for g,s in p["columns"]]
                    _write(p["question_wording"], cls, p["base_values"], p["answers"], p["data"], xl_key)
                except Exception as e:
                    errors.append((xl.sheet_names[si], str(e)))

    if progress_callback: progress_callback(1.0, "Finalizing…")

    word_bytes = excel_bytes = None
    if word_doc is not None and output_format in ("word", "both"):
        buf = io.BytesIO(); word_doc.save(buf); word_bytes = buf.getvalue()
    if xl_wb is not None and output_format in ("excel", "both"):
        if toc_entries:
            _build_xl_toc(xl_wb, toc_entries)
        buf = io.BytesIO(); xl_wb.save(buf); excel_bytes = buf.getvalue()

    return {"word_bytes": word_bytes, "excel_bytes": excel_bytes,
            "skipped": skipped, "errors": errors}


# ── Main generation function ─────────────────────────────────

def generate_outputs(
    file_bytes,
    sheet_table_configs,
    desired_groups,
    output_format,
    excel_mode,
    portrait_landscape,
    weighted_data,
    use_weighted_base=None,
    progress_callback=None,
    show_sig_flags=False,
    heatmap_scheme="None",
    heatmap_custom_start=None,
    heatmap_custom_end=None,
    fmt6_output_merged_entity=True,
    fmt6_output_t2b=True,
    fmt6_output_b2b=True,
    fmt6_output_grid=True,
    fmt6_output_mean=False,
    fmt6_output_standalone=True,
):
    if use_weighted_base is None:
        use_weighted_base = weighted_data

    xl    = pd.ExcelFile(io.BytesIO(file_bytes))
    total = max(len(sheet_table_configs), 1)

    # Route fmt6 files to dedicated handler
    fmt6_count = 0
    for cfg in sheet_table_configs[:10]:
        try:
            si  = cfg[0] if not isinstance(cfg[0], str) else xl.sheet_names.index(cfg[0])
            df  = xl.parse(si, header=None, na_values=[''])
            fmt = detect_format(df)
            if fmt == 'fmt6': fmt6_count += 1
        except Exception: pass

    if fmt6_count >= 3:
        return _generate_fmt6_outputs(
            xl, file_bytes, sheet_table_configs, desired_groups,
            output_format, excel_mode, portrait_landscape,
            weighted_data, use_weighted_base,
            fmt6_output_merged_entity, fmt6_output_t2b, fmt6_output_b2b,
            fmt6_output_grid, fmt6_output_mean, fmt6_output_standalone,
            progress_callback,
        )

    skipped      = []
    errors       = []
    parsed_cache = {}
    toc_entries  = []   # (question_wording, sheet_name)

    word_doc = None
    if output_format in ('word', 'both'):
        try: word_doc = _get_word_template(portrait_landscape)
        except Exception: word_doc = Document()
        style = word_doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)

    xl_wb         = None
    xl_sheet_ctr  = [0]
    xl_sheet_rows = {}

    if output_format in ('excel', 'both'):
        xl_wb = openpyxl.Workbook()
        xl_wb.remove(xl_wb.active)

    def _xl_get_sheet(label):
        name = label[:31]
        if excel_mode == 'per_question':
            if name not in xl_wb.sheetnames:
                xl_wb.create_sheet(title=name)
                xl_sheet_rows[name] = 1
            return xl_wb[name], xl_sheet_rows[name]
        else:
            xl_sheet_ctr[0] += 1
            sname = f"{xl_sheet_ctr[0]:03d}_{name}"[:31]
            return xl_wb.create_sheet(title=sname), 1

    def _xl_done(ws, label, next_row, question_wording=''):
        if excel_mode == 'per_question':
            xl_sheet_rows[label[:31]] = next_row
        toc_entries.append((question_wording[:100] or label, ws.title))

    def write_xl(question_wording, col_headers, base_values, answers,
                 data, multiple, sheet_label, show_base=True):
        ws, row = _xl_get_sheet(sheet_label)
        nr = _xl_write_table(ws, row, question_wording, col_headers,
                             base_values, answers, data, multiple, show_base)
        _xl_done(ws, sheet_label, nr, question_wording)

    def apply_row_filter(answers, data, row_filter):
        if row_filter == 'all' or not row_filter:
            return answers, data
        custom = {r.strip().lower() for r in row_filter}
        fa, fd = [], []
        for a, d in zip(answers, data):
            if a.strip().lower() in custom:
                fa.append(a); fd.append(d)
        return fa, fd

    first_word = True

    for idx, config_tuple in enumerate(sheet_table_configs):
        if len(config_tuple) == 4:
            sheet_idx, row_filter, group_prefix, table_label = config_tuple
        else:
            sheet_idx, row_filter = config_tuple
            group_prefix = None; table_label = None

        if progress_callback:
            progress_callback(idx / total, f"Processing sheet {sheet_idx}…")

        try:
            if isinstance(sheet_idx, str):
                sheet_name = sheet_idx
                sheet_idx  = xl.sheet_names.index(sheet_idx)
            else:
                sheet_name = xl.sheet_names[sheet_idx]
            sheet_label = sheet_name[:31]

            if sheet_idx not in parsed_cache:
                df  = xl.parse(sheet_idx, header=None, na_values=[''])
                fmt = detect_format(df)
                parsed_cache[sheet_idx] = (df, fmt)
            sheet_df, fmt = parsed_cache[sheet_idx]

            if group_prefix:
                xl_sheet_key = group_prefix[:31]
            else:
                try:
                    raw_preview  = sheet_df.values.tolist()
                    wording_row  = 3 if fmt == 'fmt5' else 2
                    wording      = str(raw_preview[wording_row][0]).strip() if (
                        len(raw_preview) > wording_row and raw_preview[wording_row][0]
                    ) else sheet_name
                    import re as _re2
                    clean = _re2.sub(r'[\\/*?\[\]:]', '', wording)[:31].strip()
                    xl_sheet_key = clean if clean else sheet_name[:31]
                except Exception:
                    xl_sheet_key = sheet_name[:31]

            # ── fmt2 ──
            if fmt == 'fmt2':
                parsed = parse_fmt2_sheet(sheet_df, desired_groups, weighted_data, weighted_base=use_weighted_base)
                if not parsed['answers']: skipped.append(sheet_label); continue
                col_labels, answers, data, flags = _prep_fmt2_fmt5(parsed, row_filter, show_sig_flags)
                if not answers: skipped.append(sheet_label); continue
                if word_doc is not None:
                    write_table_to_doc(word_doc, parsed['question_wording'], col_labels,
                                       parsed['base_values'], answers, data, 100, first_word,
                                       sig_flags=flags, heatmap_scheme=heatmap_scheme,
                                       heatmap_custom_start=heatmap_custom_start,
                                       heatmap_custom_end=heatmap_custom_end)
                    first_word = False
                if xl_wb is not None:
                    ws, row_num = _xl_get_sheet(xl_sheet_key)
                    nr = _xl_write_table(ws, row_num, parsed['question_wording'], col_labels,
                                         parsed['base_values'], answers, data, 100, True,
                                         sig_flags=flags, heatmap_scheme=heatmap_scheme,
                                         heatmap_custom_start=heatmap_custom_start,
                                         heatmap_custom_end=heatmap_custom_end)
                    _xl_done(ws, xl_sheet_key, nr, parsed['question_wording'])

            # ── fmt5 ──
            elif fmt == 'fmt5':
                parsed = parse_fmt5_sheet(sheet_df, desired_groups, weighted_data, weighted_base=use_weighted_base)
                if not parsed['answers']: skipped.append(sheet_label); continue
                col_labels, answers, data, flags = _prep_fmt2_fmt5(parsed, row_filter, show_sig_flags)
                if not answers: skipped.append(sheet_label); continue
                if word_doc is not None:
                    write_table_to_doc(word_doc, parsed['question_wording'], col_labels,
                                       parsed['base_values'], answers, data, 100, first_word,
                                       sig_flags=flags, heatmap_scheme=heatmap_scheme,
                                       heatmap_custom_start=heatmap_custom_start,
                                       heatmap_custom_end=heatmap_custom_end)
                    first_word = False
                if xl_wb is not None:
                    ws, row_num = _xl_get_sheet(xl_sheet_key)
                    nr = _xl_write_table(ws, row_num, parsed['question_wording'], col_labels,
                                         parsed['base_values'], answers, data, 100, True,
                                         sig_flags=flags, heatmap_scheme=heatmap_scheme,
                                         heatmap_custom_start=heatmap_custom_start,
                                         heatmap_custom_end=heatmap_custom_end)
                    _xl_done(ws, xl_sheet_key, nr, parsed['question_wording'])

            # ── fmt6 ──
            elif fmt == 'fmt6':
                parsed = parse_fmt6_sheet(sheet_df, desired_groups, weighted_data, weighted_base=use_weighted_base)
                if not parsed['answers']: skipped.append(sheet_label); continue
                col_labels, answers, data, flags = _prep_fmt2_fmt5(parsed, row_filter, show_sig_flags)
                if not answers: skipped.append(sheet_label); continue
                if word_doc is not None:
                    write_table_to_doc(word_doc, parsed['question_wording'], col_labels,
                                       parsed['base_values'], answers, data, 100, first_word,
                                       sig_flags=flags, heatmap_scheme=heatmap_scheme,
                                       heatmap_custom_start=heatmap_custom_start,
                                       heatmap_custom_end=heatmap_custom_end)
                    first_word = False
                if xl_wb is not None:
                    ws, row_num = _xl_get_sheet(xl_sheet_key)
                    nr = _xl_write_table(ws, row_num, parsed['question_wording'], col_labels,
                                         parsed['base_values'], answers, data, 100, True,
                                         sig_flags=flags, heatmap_scheme=heatmap_scheme,
                                         heatmap_custom_start=heatmap_custom_start,
                                         heatmap_custom_end=heatmap_custom_end)
                    _xl_done(ws, xl_sheet_key, nr, parsed['question_wording'])

            # ── fmt7 ──
            elif fmt == 'fmt7':
                parsed = parse_fmt7_sheet(sheet_df, desired_groups, weighted_data, weighted_base=use_weighted_base)
                if not parsed['answers']: skipped.append(sheet_label); continue
                col_labels, answers, data, flags = _prep_fmt2_fmt5(parsed, row_filter, show_sig_flags)
                if not answers: skipped.append(sheet_label); continue
                if word_doc is not None:
                    write_table_to_doc(word_doc, parsed['question_wording'], col_labels,
                                       parsed['base_values'], answers, data, 100, first_word,
                                       sig_flags=flags, heatmap_scheme=heatmap_scheme,
                                       heatmap_custom_start=heatmap_custom_start,
                                       heatmap_custom_end=heatmap_custom_end)
                    first_word = False
                if xl_wb is not None:
                    ws, row_num = _xl_get_sheet(xl_sheet_key)
                    nr = _xl_write_table(ws, row_num, parsed['question_wording'], col_labels,
                                         parsed['base_values'], answers, data, 100, True,
                                         sig_flags=flags, heatmap_scheme=heatmap_scheme,
                                         heatmap_custom_start=heatmap_custom_start,
                                         heatmap_custom_end=heatmap_custom_end)
                    _xl_done(ws, xl_sheet_key, nr, parsed['question_wording'])

            # ── fmt3 ──
            elif fmt == 'fmt3':
                parsed     = parse_fmt3_sheet(sheet_df, desired_groups)
                if not parsed['answers']: skipped.append(sheet_label); continue
                col_labels    = [g for (g, s) in parsed['columns']]
                companies     = parsed['answers']
                company_bases = parsed['company_bases']
                pct_data      = parsed['data']
                n_cols        = len(col_labels)

                if row_filter not in ('all', None):
                    filtered = [(c, d, b) for c, d, b in zip(companies, pct_data, company_bases)
                                if row_filter == 'all' or c.strip().lower() in {r.lower() for r in (row_filter if isinstance(row_filter, list) else [])}]
                    companies, pct_data, company_bases = (list(x) for x in zip(*filtered)) if filtered else ([], [], [])

                if not companies: skipped.append(sheet_label); continue

                def _write_fmt3_word(show_n, is_first_arg, suffix=''):
                    if word_doc is None: return
                    if not is_first_arg: word_doc.add_paragraph()
                    qp = word_doc.add_paragraph()
                    qp.style = word_doc.styles['Normal']
                    qp.text  = parsed['question_wording'] + suffix
                    tbl = word_doc.add_table(rows=1, cols=n_cols + 1)
                    tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    hdr = tbl.rows[0].cells; hdr[0].text = ''
                    for ci in range(n_cols):
                        hdr[ci+1].text = ''
                        hdr[ci+1].paragraphs[0].add_run(col_labels[ci]).bold = True
                        hdr[ci+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for ri, company in enumerate(companies):
                        rc = tbl.add_row().cells; rc[0].text = company
                        bases = company_bases[ri] if ri < len(company_bases) else []
                        pcts  = pct_data[ri]      if ri < len(pct_data)      else []
                        for ci in range(n_cols):
                            pct  = pcts[ci]  if ci < len(pcts)  else None
                            n    = bases[ci] if ci < len(bases) else None
                            cell = rc[ci + 1]; cell.text = ''
                            if pct is None or isinstance(pct, str):
                                cell.paragraphs[0].add_run('-')
                            else:
                                pct_str = f"{normal_round(round(pct * 100, 3))}%"
                                n_str   = f"\n(N={prettyPrint(n)})" if (show_n and n) else ''
                                cell.paragraphs[0].add_run(pct_str + n_str)
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for ci in range(n_cols + 1):
                        for cell in tbl.columns[ci].cells:
                            cell.width = Inches(3) if ci == 0 else Inches(1.75)

                _write_fmt3_word(show_n=False, is_first_arg=first_word)
                _write_fmt3_word(show_n=True,  is_first_arg=False, suffix=' (with base)')
                first_word = False

                if xl_wb is not None:
                    write_xl(parsed['question_wording'], col_labels, [None]*n_cols,
                             companies, pct_data, 100, xl_sheet_key, show_base=False)
                    data_with_n = []
                    for ri, company in enumerate(companies):
                        bases = company_bases[ri] if ri < len(company_bases) else []
                        pcts  = pct_data[ri]
                        row   = []
                        for ci in range(n_cols):
                            pct = pcts[ci]  if ci < len(pcts)  else None
                            n   = bases[ci] if ci < len(bases) else None
                            if pct is None or isinstance(pct, str):
                                row.append(None)
                            else:
                                row.append(f"{normal_round(round(pct*100,3))}%" + (f" (N={prettyPrint(n)})" if n else ''))
                        data_with_n.append(row)
                    write_xl(parsed['question_wording'] + ' (with base)',
                             col_labels, [None]*n_cols, companies, data_with_n, 1, xl_sheet_key, show_base=False)

            # ── fmt4 ──
            elif fmt == 'fmt4':
                parsed  = parse_fmt4_sheet(sheet_df)
                if not parsed['answers']: skipped.append(sheet_label); continue
                answers, data = apply_row_filter(parsed["answers"], parsed["data"], row_filter)
                if not answers: skipped.append(sheet_label); continue
                brands      = parsed['brands']
                base_values = parsed['base_values']
                n_brands    = len(brands)

                if word_doc is not None:
                    if not first_word: word_doc.add_paragraph()
                    qp = word_doc.add_paragraph()
                    qp.style = word_doc.styles['Normal']
                    qp.text  = parsed['question_wording']
                    tbl = word_doc.add_table(rows=1, cols=n_brands + 1)
                    tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    hdr = tbl.rows[0].cells; hdr[0].text = ''
                    for ci, brand in enumerate(brands):
                        n = base_values[ci] if ci < len(base_values) else None
                        hdr[ci+1].text = ''
                        hdr[ci+1].paragraphs[0].add_run(f"{brand}\n(N={prettyPrint(n)})").bold = True
                        hdr[ci+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for ri, answer in enumerate(answers):
                        rc = tbl.add_row().cells; rc[0].text = answer
                        row_vals = data[ri] if ri < len(data) else []
                        for ci in range(n_brands):
                            val  = row_vals[ci] if ci < len(row_vals) else None
                            cell = rc[ci + 1]; cell.text = ''
                            if val is None: cell.paragraphs[0].add_run('-')
                            else: cell.paragraphs[0].add_run(f"{normal_round(round(val*100,3))}%")
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for ci in range(n_brands + 1):
                        for cell in tbl.columns[ci].cells:
                            cell.width = Inches(2) if ci == 0 else Inches(1.25)
                    first_word = False

                if xl_wb is not None:
                    write_xl(parsed['question_wording'], brands, base_values, answers, data, 100, xl_sheet_key)

            else:
                skipped.append(f"{sheet_label} (fmt1/grid — skipped)")

        except Exception as e:
            errors.append((sheet_name if 'sheet_name' in dir() else str(sheet_idx), str(e)))

    if progress_callback:
        progress_callback(1.0, "Finalizing…")

    word_bytes  = None
    excel_bytes = None

    if word_doc is not None and output_format in ('word', 'both'):
        buf = io.BytesIO(); word_doc.save(buf); word_bytes = buf.getvalue()

    if xl_wb is not None and output_format in ('excel', 'both'):
        if toc_entries:
            _build_xl_toc(xl_wb, toc_entries)
        buf = io.BytesIO(); xl_wb.save(buf); excel_bytes = buf.getvalue()

    return {
        'word_bytes':  word_bytes,
        'excel_bytes': excel_bytes,
        'skipped':     skipped,
        'errors':      errors,
    }


# ── Multi-source scanning ─────────────────────────────────────

def scan_multi_source(file_list):
    import re
    banners = []
    for file_bytes, banner_name in file_list:
        metas    = scan_file(file_bytes)
        all_cols = get_all_columns(metas)
        banners.append({'name': banner_name, 'sheets': metas,
                        'all_cols': all_cols, 'file_bytes': file_bytes})

    if not banners:
        return {'banners': [], 'matched': [], 'unmatched': []}

    def q_id(wording):
        m = re.match(r'^([A-Za-z]+\d+)', wording.strip())
        return m.group(1).upper() if m else None

    banner_indices = []
    for b in banners:
        idx = {}
        for m in b['sheets']:
            if m['fmt'] == 'error': continue
            qid = q_id(m['question_wording'])
            key = qid if qid else m['question_wording'].strip()[:80]
            if key in idx: key = f"{key}_{m['index']}"
            idx[key] = m
        banner_indices.append(idx)

    ref_idx  = banner_indices[0]
    matched  = []
    unmatched = []

    for key, ref_meta in ref_idx.items():
        entry = {'q_id': key, 'wording': ref_meta['question_wording'],
                 'banner_sheets': {banners[0]['name']: ref_meta}}
        for bi in range(1, len(banners)):
            other_idx = banner_indices[bi]
            entry['banner_sheets'][banners[bi]['name']] = other_idx.get(key)
        matched.append(entry)

    for bi in range(1, len(banners)):
        other_idx = banner_indices[bi]
        for key, meta in other_idx.items():
            if key not in ref_idx:
                unmatched.append({'banner_name': banners[bi]['name'],
                                  'q_id': key, 'sheet_meta': meta})

    return {'banners': banners, 'matched': matched, 'unmatched': unmatched}


def generate_merged_outputs(
    multi_source, matched_overrides, selected_cols_per_banner,
    output_format, excel_mode, portrait_landscape,
    weighted_data, row_filter, progress_callback=None,
):
    banners  = multi_source['banners']
    matched  = multi_source['matched']
    skipped  = []
    errors   = []

    effective_matched = []
    for entry in matched:
        eff = dict(entry)
        if entry['q_id'] in (matched_overrides or {}):
            overrides = matched_overrides[entry['q_id']]
            for bn, sheet_idx in overrides.items():
                if sheet_idx is None:
                    eff['banner_sheets'][bn] = None
                else:
                    for b in banners:
                        if b['name'] == bn:
                            for m in b['sheets']:
                                if m['index'] == sheet_idx:
                                    eff['banner_sheets'][bn] = m
        effective_matched.append(eff)

    for entry in multi_source.get('unmatched', []):
        qid = entry['q_id']
        if qid in (matched_overrides or {}):
            overrides = matched_overrides[qid]
            new_entry = {'q_id': qid, 'wording': entry['sheet_meta']['question_wording'],
                         'banner_sheets': {}}
            for b in banners: new_entry['banner_sheets'][b['name']] = None
            new_entry['banner_sheets'][entry['banner_name']] = entry['sheet_meta']
            for bn, sheet_idx in overrides.items():
                if sheet_idx is not None:
                    for b in banners:
                        if b['name'] == bn:
                            for m in b['sheets']:
                                if m['index'] == sheet_idx:
                                    new_entry['banner_sheets'][bn] = m
            effective_matched.append(new_entry)

    word_doc = None
    if output_format in ('word', 'both'):
        try: word_doc = _get_word_template(portrait_landscape)
        except Exception: word_doc = Document()
        style = word_doc.styles['Normal']
        style.font.name = 'Arial'; style.font.size = Pt(10)

    xl_wb         = None
    xl_sheet_ctr  = [0]
    xl_sheet_rows = {}
    toc_entries   = []

    if output_format in ('excel', 'both'):
        xl_wb = openpyxl.Workbook()
        xl_wb.remove(xl_wb.active)

    def _xl_get_sheet(label):
        name = label[:31]
        if excel_mode == 'per_question':
            if name not in xl_wb.sheetnames:
                xl_wb.create_sheet(title=name)
                xl_sheet_rows[name] = 1
            return xl_wb[name], xl_sheet_rows.get(name, 1)
        else:
            xl_sheet_ctr[0] += 1
            sname = f"{xl_sheet_ctr[0]:03d}_{name}"[:31]
            return xl_wb.create_sheet(title=sname), 1

    def _xl_done(ws, label, next_row, q_wording=''):
        if excel_mode == 'per_question':
            xl_sheet_rows[label[:31]] = next_row
        toc_entries.append((q_wording[:100] or label, ws.title))

    def _apply_row_filter(answers, data):
        if row_filter == 'all' or not row_filter: return answers, data
        custom = {r.strip().lower() for r in row_filter}
        fa, fd = [], []
        for a, d in zip(answers, data):
            if a.strip().lower() in custom: fa.append(a); fd.append(d)
        return fa, fd

    total     = max(len(effective_matched), 1)
    first_doc = True
    bytes_map = {b['name']: b['file_bytes'] for b in banners}

    for idx, entry in enumerate(effective_matched):
        if progress_callback: progress_callback(idx / total, f"Processing {entry['q_id']}…")
        try:
            q_id_key = entry['q_id']
            xl_key   = q_id_key[:31]
            merged_col_labels  = []
            merged_base_values = []
            merged_answers     = None
            merged_data_cols   = []

            for b in banners:
                bname      = b['name']
                sheet_meta = entry['banner_sheets'].get(bname)
                sel_cols   = selected_cols_per_banner.get(bname, [])
                if sheet_meta is None or not sel_cols:
                    for col in sel_cols:
                        merged_col_labels.append(col)
                        merged_base_values.append(None)
                        merged_data_cols.append(None)
                    continue
                fb  = bytes_map[bname]
                xl2 = pd.ExcelFile(io.BytesIO(fb))
                df  = xl2.parse(sheet_meta['index'], header=None, na_values=[''])
                fmt = detect_format(df)
                if fmt == 'fmt2':   parsed = parse_fmt2_sheet(df, sel_cols, weighted_data)
                elif fmt == 'fmt5': parsed = parse_fmt5_sheet(df, sel_cols, weighted_data)
                else:
                    skipped.append(f"{bname} — {sheet_meta['sheet_name']} (fmt not supported in merge)")
                    for col in sel_cols:
                        merged_col_labels.append(col)
                        merged_base_values.append(None)
                        merged_data_cols.append(None)
                    continue

                col_names = [g for (g, s) in parsed['columns']]
                answers_b = parsed['answers']
                data_b    = parsed['data']
                if merged_answers is None: merged_answers = answers_b
                ans_lookup = {ans.strip().lower(): data_b[ai] for ai, ans in enumerate(answers_b)}

                for ci, col_name in enumerate(col_names):
                    merged_col_labels.append(col_name)
                    merged_base_values.append(parsed['base_values'][ci] if ci < len(parsed['base_values']) else None)
                    col_data = []
                    for ans in (merged_answers or []):
                        row = ans_lookup.get(ans.strip().lower())
                        col_data.append(row[ci] if (row and ci < len(row)) else None)
                    merged_data_cols.append(col_data)

            if not merged_answers: skipped.append(q_id_key); continue
            n_answers = len(merged_answers)
            for ci in range(len(merged_data_cols)):
                if merged_data_cols[ci] is None:
                    merged_data_cols[ci] = [None] * n_answers

            merged_data = []
            for ai in range(n_answers):
                row = [merged_data_cols[ci][ai] if ai < len(merged_data_cols[ci]) else None
                       for ci in range(len(merged_col_labels))]
                merged_data.append(row)

            filtered_answers, filtered_data = _apply_row_filter(merged_answers, merged_data)
            if not filtered_answers: skipped.append(q_id_key); continue

            question_wording = entry['wording']

            if word_doc is not None:
                write_table_to_doc(word_doc, question_wording, merged_col_labels,
                                   merged_base_values, filtered_answers, filtered_data, 100, first_doc)
                first_doc = False

            if xl_wb is not None:
                ws, row_num = _xl_get_sheet(xl_key)
                next_row = _xl_write_table(ws, row_num, question_wording, merged_col_labels,
                                           merged_base_values, filtered_answers, filtered_data, 100, True)
                _xl_done(ws, xl_key, next_row, question_wording)

        except Exception as e:
            errors.append((entry.get('wording', entry['q_id'])[:60], str(e)))

    if progress_callback: progress_callback(1.0, "Finalizing…")

    word_bytes = excel_bytes = None
    if word_doc is not None and output_format in ('word', 'both'):
        buf = io.BytesIO(); word_doc.save(buf); word_bytes = buf.getvalue()
    if xl_wb is not None and output_format in ('excel', 'both'):
        if toc_entries: _build_xl_toc(xl_wb, toc_entries)
        buf = io.BytesIO(); xl_wb.save(buf); excel_bytes = buf.getvalue()

    return {'word_bytes': word_bytes, 'excel_bytes': excel_bytes,
            'skipped': skipped, 'errors': errors}
