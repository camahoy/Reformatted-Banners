"""
engine.py — Banner Formatter core logic v2.7
All parsing, detection, and output writing lives here.
The Streamlit app (app.py) calls these functions directly.
"""

print("ENGINE VERSION 2.7 LOADED")

import io
import math
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import openpyxl
from openpyxl.styles import Font as XLFont, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter


# ── Helper functions ─────────────────────────────────────────

def lowerList(inputList):
    return [x.lower() if isinstance(x, str) else x for x in inputList]

def normal_round(n):
    if n - math.floor(n) < 0.5:
        return math.floor(n)
    return math.ceil(n)

def prettyPrint(number):
    if isinstance(number, str):
        return number
    try:
        if np.isnan(number):
            return ""
    except Exception:
        pass
    if isinstance(number, (int, float)):
        newNumber = str(normal_round(number))
        if len(newNumber) > 3:
            newNumber = newNumber[:-3] + ',' + newNumber[-3:]
        return newNumber
    return number

def removeNaN(theList):
    return [x for x in theList if not (isinstance(x, float) and math.isnan(x))]

def transposeList(theList):
    return pd.DataFrame(theList).T.values.tolist()

def sum2(listSum):
    total = 0
    for i in listSum:
        try:
            total += float(i)
        except Exception:
            pass
    return total


# ── Format detection ─────────────────────────────────────────

def detect_format(sheet_df):
    """
    Returns 'fmt2', 'fmt3', 'fmt4', or 'fmt1' for each sheet.
    fmt2 = standard question (countries as columns)
    fmt3 = top box summary (companies as rows, countries as columns, floating base)
    fmt4 = grid table (brands as columns, scale options as rows)
    fmt1 = original transposed format (legacy / per-country)
    """
    raw = sheet_df.values.tolist()

    # fmt6: like fmt5 but with extra descriptor row at position 2
    # Signature: row 2 col 0 is a short label (not a question), row 3 col 0 is question,
    #            row 4 col 1 = 'Total', row 5 col 1 = category name
    try:
        row2_col0 = raw[2][0]
        row3_col0 = raw[3][0]
        row4_col1 = raw[4][1]
        row5_col1 = raw[5][1]
        if (
            isinstance(row3_col0, str) and len(row3_col0.strip()) > 5
            and isinstance(row4_col1, str) and 'total' in row4_col1.lower()
            and isinstance(row5_col1, str) and row5_col1.strip()
            and (not isinstance(row2_col0, str) or len(row2_col0.strip()) < 50)
        ):
            return 'fmt6'
    except (IndexError, TypeError):
        pass

    # fmt5: question at row 3, categories at row 5, two base rows (unweighted + weighted)
    # Signature: row 3 col 0 has question wording, row 4 col 1 has 'Total',
    #            row 5 col 1 has category names
    try:
        row3_col0 = raw[3][0]
        row4_col1 = raw[4][1]
        row5_col1 = raw[5][1]
        if (
            isinstance(row3_col0, str) and len(row3_col0.strip()) > 5
            and isinstance(row4_col1, str) and 'total' in row4_col1.lower()
            and isinstance(row5_col1, str) and row5_col1.strip()
        ):
            return 'fmt5'
    except (IndexError, TypeError):
        pass

    # fmt4: brands as columns — row 3 col 1 blank, row 4 col 1 is a brand name not 'Total'
    try:
        row3_col1 = raw[3][1]
        row4_col1 = raw[4][1]
        if (
            isinstance(row4_col1, str) and row4_col1.strip()
            and 'total' not in str(row4_col1).lower()
            and (
                row3_col1 is None
                or not isinstance(row3_col1, str)
                or row3_col1.strip() == ''
                or (isinstance(row3_col1, float) and math.isnan(row3_col1))
            )
        ):
            return 'fmt4'
    except (IndexError, TypeError):
        pass

    # fmt7: weighted banner with group header row
    # Row 2: question wording
    # Row 3: group names (Gender, Generation...) — col 1 may be blank
    # Row 4: category names (Total, Male, Female...) — col 1 = 'Total'
    # Row 5: letter codes
    # Row 7: Unweighted Base + counts
    # Row 9: Base: Total Answering + weighted counts
    try:
        row2_col0 = raw[2][0]
        row4_col1 = raw[4][1]
        row7_col0 = raw[7][0] if len(raw) > 7 else None
        if (
            isinstance(row2_col0, str) and len(row2_col0.strip()) > 10
            and isinstance(row4_col1, str) and 'total' in row4_col1.lower()
            and isinstance(row7_col0, str) and row7_col0.strip().lower().startswith('unweighted')
        ):
            return 'fmt7'
    except (IndexError, TypeError):
        pass

    # fmt7: extra group header row between question and categories
    # Row 2=question, Row 3=group names, Row 4=categories (col1=Total), Row 7=UnweightedBase
    try:
        row2_col0 = raw[2][0]
        row4_col1 = raw[4][1]
        row7_col0 = raw[7][0] if len(raw) > 7 else None
        if (isinstance(row2_col0, str) and len(row2_col0.strip()) > 10
            and isinstance(row4_col1, str) and 'total' in row4_col1.lower()
            and isinstance(row7_col0, str) and row7_col0.strip().lower().startswith('unweighted')):
            return 'fmt7'
    except (IndexError, TypeError):
        pass

    # fmt2 / fmt3: row 3 col 1 = 'Total', row 4 = sub-labels
    try:
        row2_col0 = raw[2][0]
        row3_col1 = raw[3][1]
        if (
            isinstance(row2_col0, str) and len(row2_col0.strip()) > 10
            and isinstance(row3_col1, str) and 'total' in row3_col1.lower()
        ):
            # Check rows 7-11 for a base label (must start with 'base' or 'unweighted')
            base_row_col0 = None
            base_row_col1 = None
            for check_row in range(7, min(12, len(raw))):
                cell = raw[check_row][0] if raw[check_row] else None
                if isinstance(cell, str):
                    stripped = cell.strip().lower()
                    if stripped.startswith('base') or stripped.startswith('unweighted'):
                        base_row_col0 = cell
                        base_row_col1 = raw[check_row][1] if len(raw[check_row]) > 1 else None
                        break

            if base_row_col0 is not None:
                if base_row_col1 is None or (isinstance(base_row_col1, float) and math.isnan(base_row_col1)):
                    return 'fmt3'
                return 'fmt2'
    except (IndexError, TypeError):
        pass

    return 'fmt1'


# ── Sheet scanning ───────────────────────────────────────────

def scan_file(file_bytes):
    """
    Read an uploaded Excel file and return metadata for every sheet.
    Returns list of dicts: {index, name, fmt, question_wording, columns}
    """
    xl = pd.ExcelFile(io.BytesIO(file_bytes))
    results = []
    for i, sheet_name in enumerate(xl.sheet_names):
        # Skip sheet 0 only if it looks like a table of contents
        # (name is 'Index', 'TOC', 'Contents', or has only 1 column of data)
        if i == 0:
            try:
                df0 = xl.parse(0, header=None, na_values=[''])
                fmt0 = detect_format(df0)
                # If sheet 0 is a real data format, include it
                if fmt0 == 'fmt1':
                    continue  # likely TOC/index
                # Otherwise fall through and process it
            except Exception:
                continue
        try:
            df  = xl.parse(i, header=None, na_values=[''])
            fmt = detect_format(df)
            raw = df.values.tolist()
            # Question wording: fmt5/fmt6 have it at row 3, others at row 2
            if fmt in ('fmt5', 'fmt6'):
                question_wording = str(raw[3][0]).strip() if (
                    len(raw) > 3 and raw[3] and raw[3][0]
                ) else sheet_name
            else:
                question_wording = str(raw[2][0]).strip() if (
                    len(raw) > 2 and raw[2] and raw[2][0]
                ) else sheet_name

            # Collect available columns
            columns = []
            if fmt in ('fmt2', 'fmt3'):
                group_row    = raw[3] if len(raw) > 3 else []
                sublabel_row = raw[4] if len(raw) > 4 else []
                for j in range(1, len(group_row)):
                    g = group_row[j]
                    s = sublabel_row[j] if j < len(sublabel_row) else ''
                    if isinstance(g, str) and g.strip():
                        columns.append((g.strip(), s.strip() if isinstance(s, str) else ''))
            elif fmt == 'fmt5':
                cat_row = raw[5] if len(raw) > 5 else []
                for j in range(1, len(cat_row)):
                    g = cat_row[j]
                    if isinstance(g, str) and g.strip() and g.strip() != '\xa0':
                        columns.append((g.strip(), ''))
            elif fmt == 'fmt6':
                cat_row = raw[5] if len(raw) > 5 else []
                for j in range(1, len(cat_row)):
                    g = cat_row[j]
                    if isinstance(g, str) and g.strip() and g.strip() != '\xa0':
                        columns.append((g.strip(), ''))
            elif fmt == 'fmt7':
                cat_row = raw[4] if len(raw) > 4 else []
                for j in range(1, len(cat_row)):
                    g = cat_row[j]
                    if isinstance(g, str) and g.strip() and g.strip() != '\xa0':
                        columns.append((g.strip(), ''))
            elif fmt == 'fmt7':
                cat_row = raw[4] if len(raw) > 4 else []
                for j in range(1, len(cat_row)):
                    g = cat_row[j]
                    if isinstance(g, str) and g.strip() and g.strip() != '\xa0':
                        columns.append((g.strip(), ''))
            elif fmt == 'fmt4':
                brand_row = raw[4] if len(raw) > 4 else []
                for j in range(1, len(brand_row)):
                    v = brand_row[j]
                    if isinstance(v, str) and v.strip():
                        columns.append((v.strip(), ''))

            results.append({
                'index':            i,
                'sheet_name':       sheet_name,
                'fmt':              fmt,
                'question_wording': question_wording,
                'columns':          columns,
            })
        except Exception as e:
            results.append({
                'index':            i,
                'sheet_name':       sheet_name,
                'fmt':              'error',
                'question_wording': f'Error reading sheet: {e}',
                'columns':          [],
            })
    return results


def get_all_columns(sheet_metas):
    """
    Return deduplicated list of (group, sublabel) column tuples
    from fmt2/fmt3/fmt5 sheets (not fmt4 — those use brand columns).
    """
    seen = set()
    cols = []
    for s in sheet_metas:
        if s['fmt'] in ('fmt2', 'fmt3', 'fmt5', 'fmt6', 'fmt7'):
            for col in s['columns']:
                if col[0] not in seen:
                    seen.add(col[0])
                    cols.append(col)
    return cols


def scan_rows_for_sheets(file_bytes, sheet_indices):
    """
    Return deduplicated ordered list of answer/row labels
    across all given sheet indices. Used to populate row checkboxes.
    """
    xl   = pd.ExcelFile(io.BytesIO(file_bytes))
    seen = set()
    rows = []
    for i in sheet_indices:
        try:
            df  = xl.parse(i, header=None, na_values=[''])
            fmt = detect_format(df)
            raw = df.values.tolist()
            if fmt == 'fmt2':
                # base row index
                base_row_idx = None
                for ri, row in enumerate(raw):
                    if isinstance(row[0], str) and row[0].strip().lower().startswith('base'):
                        base_row_idx = ri
                        break
                start = (base_row_idx + 2) if base_row_idx is not None else 9
                j = start
                while j < len(raw):
                    label = raw[j][0]
                    if isinstance(label, str) and label.strip():
                        clean = label.strip()
                        if clean.lower() == 'sigma':
                            break
                        if clean not in seen:
                            seen.add(clean)
                            rows.append(clean)
                    j += 3
            elif fmt in ('fmt5', 'fmt6', 'fmt7'):
                # fmt5/fmt6/fmt7 data starts at row 11-12
                j = 11
                while j < len(raw):
                    label = raw[j][0]
                    if isinstance(label, str) and label.strip():
                        clean = label.strip()
                        if clean.lower() == 'sigma':
                            break
                        if clean not in seen:
                            seen.add(clean)
                            rows.append(clean)
                    j += 3
            elif fmt == 'fmt3':
                end_markers = {'overlap formula used', 'sigma'}
                j = 8
                while j < len(raw):
                    label = raw[j][0]
                    if isinstance(label, str) and label.strip():
                        clean = label.strip()
                        if clean.lower() in end_markers:
                            break
                        if clean not in seen:
                            seen.add(clean)
                            rows.append(clean)
                    j += 3
            elif fmt == 'fmt4':
                j = 9
                while j < len(raw):
                    label = raw[j][0]
                    if isinstance(label, str) and label.strip():
                        clean = label.strip()
                        if clean.lower() == 'sigma':
                            break
                        if clean not in seen:
                            seen.add(clean)
                            rows.append(clean)
                    j += 3
        except Exception:
            pass
    return rows


def get_question_groups(sheet_metas):
    """
    Group sheets by question prefix (A0, A1, S1 etc).
    Returns ordered list of dicts:
      { prefix, sheets: [meta, ...], fmt }
    fmt is taken from first sheet in group.
    """
    import re
    groups = {}
    order  = []
    for m in sheet_metas:
        if m['fmt'] == 'error':
            continue
        # Extract prefix: leading letters + digits e.g. "A0", "S1", "A10"
        match = re.match(r'^([A-Za-z]+\d+)', m['question_wording'].strip())
        if match:
            prefix = match.group(1).upper()
        else:
            prefix = m['sheet_name'][:6]
        if prefix not in groups:
            groups[prefix] = {'prefix': prefix, 'sheets': [], 'fmt': m['fmt']}
            order.append(prefix)
        groups[prefix]['sheets'].append(m)
    return [groups[p] for p in order]


# ── Parsers ──────────────────────────────────────────────────

def _select_cols(all_cols, desired_groups, sublabel_filter='total'):
    """Filter column list to desired_groups (list of group name strings)."""
    if desired_groups is not None:
        desired_lower = {d.lower() for d in desired_groups}
        return [(j, g, s) for (j, g, s) in all_cols if g.lower() in desired_lower]
    # Auto: take columns whose sub_label is 'Total' + first column
    selected = [(j, g, s) for (j, g, s) in all_cols if s.lower() == sublabel_filter]
    if all_cols and all_cols[0] not in selected:
        selected.insert(0, all_cols[0])
    return selected


def parse_fmt2_sheet(sheet_df, desired_groups=None, weighted_data=False, weighted_base=None):
    """
    weighted_data: affects data_start (how many base rows to skip before answers)
    weighted_base: which base row to use for N= display
                   None = same as weighted_data
                   True = use weighted base (row base_row_idx + 2)
                   False = use unweighted base (row base_row_idx + 0)
    """
    if weighted_base is None:
        weighted_base = weighted_data
    raw = sheet_df.values.tolist()
    question_wording = str(raw[2][0]).strip() if raw[2][0] else ''

    group_row    = raw[3]
    sublabel_row = raw[4]
    all_cols     = []
    for j in range(1, len(group_row)):
        g = group_row[j]
        s = sublabel_row[j]
        if isinstance(g, str) and g.strip():
            all_cols.append((j, g.strip(), s.strip() if isinstance(s, str) else ''))

    selected_cols = _select_cols(all_cols, desired_groups)
    col_indices   = [j for (j, g, s) in selected_cols]
    col_labels    = [(g, s) for (j, g, s) in selected_cols]

    # Base row — find first row where col 0 is a base label
    # Must start with 'base' OR be 'unweighted base' — never match mid-sentence 'based'
    BASE_LABELS = {'base', 'unweighted base', 'weighted base'}
    base_row_idx = None
    for i, row in enumerate(raw):
        cell = row[0] if row else None
        if isinstance(cell, str):
            stripped = cell.strip().lower()
            # Match if it starts with 'base' or 'unweighted' followed by 'base'
            if (stripped.startswith('base') or
                stripped.startswith('unweighted') or
                stripped.startswith('weighted base')):
                base_row_idx = i
                break

    # Base display: use weighted_base to pick which row to show in N=
    # weighted_base=True  → weighted count (row +2 from unweighted base row)
    # weighted_base=False → unweighted count (row +0, same row as label)
    base_offset   = 2 if weighted_base else 0
    base_values   = []
    if base_row_idx is not None:
        base_data_row = raw[base_row_idx + base_offset] if base_row_idx + base_offset < len(raw) else []
        base_values   = [base_data_row[j] if j < len(base_data_row) else None for j in col_indices]

    # Data start: skip both base rows when weighted_data=True
    # weighted_data=True:  unweighted(+0) + blank(+1) + weighted(+2) + blank(+3) = data at +4
    # weighted_data=False: base(+0) + blank(+1) = data at +2
    base_rows_used = 4 if weighted_data else 2
    data_start = (base_row_idx + base_rows_used) if base_row_idx is not None else 9

    # Answers + data
    answers  = []
    data     = []
    sig_data = []

    # Column letter codes (row 5, 0-based) — used for sig flag logic
    letter_row  = raw[5] if len(raw) > 5 else []
    col_letters = [letter_row[j] if j < len(letter_row) else None for j in col_indices]

    i = data_start
    while i < len(raw):
        label = raw[i][0]
        if isinstance(label, str) and label.strip():
            label_clean = label.strip()
            if label_clean.lower() == 'sigma':
                break
            answers.append(label_clean)
            pct_row = raw[i + 1] if i + 1 < len(raw) else []
            sig_row = raw[i + 2] if i + 2 < len(raw) else []
            row_vals = [pct_row[j] if j < len(pct_row) else None for j in col_indices]
            sig_vals = [sig_row[j]  if j < len(sig_row)  else None for j in col_indices]
            data.append(row_vals)
            sig_data.append(sig_vals)
        i += 3

    return {
        'question_wording': question_wording,
        'columns':          col_labels,
        'base_values':      base_values,
        'answers':          answers,
        'data':             data,
        'sig_data':         sig_data,
        'col_letters':      col_letters,
    }


def parse_fmt3_sheet(sheet_df, desired_groups=None):
    """
    Handles two confirmed fmt3 layouts:

    Unweighted (original):
      Row 7:  Base label + counts (same row)
      Row 8:  blank
      Row 9+: [company + counts] [%] [sig]  every 3 rows

    Weighted:
      Row 7:  Unweighted Base + counts
      Row 8:  blank
      Row 9:  Base: ... label (no counts on this row)
      Row 10: blank
      Row 11: counts (no label)
      Row 12: company label + % values
      Row 13: sig letters
      Pattern: [counts] [label+%] [sig]  every 3 rows, starting at row 11
    """
    raw = sheet_df.values.tolist()
    question_wording = str(raw[2][0]).strip() if raw[2][0] else ''

    group_row    = raw[3]
    sublabel_row = raw[4]
    all_cols     = []
    for j in range(1, len(group_row)):
        g = group_row[j]
        s = sublabel_row[j]
        if isinstance(g, str) and g.strip():
            all_cols.append((j, g.strip(), s.strip() if isinstance(s, str) else ''))

    selected_cols = _select_cols(all_cols, desired_groups)
    col_indices   = [j for (j, g, s) in selected_cols]
    col_labels    = [(g, s) for (j, g, s) in selected_cols]

    end_markers = {'overlap formula used', 'sigma'}

    def get_val(row, j):
        v = row[j] if j < len(row) else None
        if v is None:
            return None
        if isinstance(v, str) and v.strip() in ('-', '', '\xa0'):
            return None
        if isinstance(v, float) and math.isnan(v):
            return None
        return v

    def is_numeric(v):
        return isinstance(v, (int, float)) and not (isinstance(v, float) and math.isnan(v))

    def row_has_nums(row):
        return any(is_numeric(row[j]) for j in col_indices if j < len(row))

    def row_has_label(row):
        c = row[0] if row else None
        return isinstance(c, str) and c.strip()

    # Detect layout by scanning from row 7 onwards
    # Find first row that has numeric data in the data columns
    # Then check if the NEXT row has a company label (weighted) or
    # if THIS row has a company label (unweighted)
    weighted_fmt3 = False
    data_start    = 8

    for check in range(7, min(16, len(raw))):
        row = raw[check]
        label = row[0] if row else None
        # Skip base/blank rows
        if not isinstance(label, str) or not label.strip():
            continue
        stripped = label.strip().lower()
        if stripped.startswith('base') or stripped.startswith('unweighted'):
            continue
        if stripped in end_markers:
            break
        # This row has a company label — check if PREVIOUS row had numbers
        prev = raw[check - 1] if check > 0 else []
        if row_has_nums(prev):
            # Weighted: counts are on prev row, label+% on this row
            weighted_fmt3 = True
            data_start    = check - 1
        else:
            # Unweighted: label+counts on this row
            weighted_fmt3 = False
            data_start    = check
        break

    companies     = []
    company_bases = []
    data          = []

    i = data_start
    while i < len(raw):
        if weighted_fmt3:
            # [counts row i] [label+% row i+1] [sig row i+2]
            counts_row = raw[i]
            label_row  = raw[i + 1] if i + 1 < len(raw) else []
            label      = label_row[0] if label_row else None
            if not isinstance(label, str) or not label.strip():
                i += 1
                continue
            label_clean = label.strip()
            if label_clean.lower() in end_markers:
                break
            companies.append(label_clean)
            company_bases.append([get_val(counts_row, j) for j in col_indices])
            data.append([get_val(label_row, j) for j in col_indices])
            i += 3
        else:
            # [label+counts row i] [% row i+1] [sig row i+2]
            label_row = raw[i]
            label     = label_row[0] if label_row else None
            if not isinstance(label, str) or not label.strip():
                i += 1
                continue
            label_clean = label.strip()
            if label_clean.lower() in end_markers:
                break
            pct_row = raw[i + 1] if i + 1 < len(raw) else []
            companies.append(label_clean)
            company_bases.append([get_val(label_row, j) for j in col_indices])
            data.append([get_val(pct_row, j) for j in col_indices])
            i += 3

    return {
        'question_wording': question_wording,
        'columns':          col_labels,
        'base_values':      None,
        'company_bases':    company_bases,
        'answers':          companies,
        'data':             data,
    }



def parse_fmt4_sheet(sheet_df):
    raw = sheet_df.values.tolist()
    question_wording = str(raw[2][0]).strip() if raw[2][0] else ''

    brand_row     = raw[4]
    brands        = []
    brand_indices = []
    for j in range(1, len(brand_row)):
        v = brand_row[j]
        if isinstance(v, str) and v.strip():
            brands.append(v.strip())
            brand_indices.append(j)

    base_row    = raw[7]
    base_values = [base_row[j] if j < len(base_row) else None for j in brand_indices]

    answers = []
    data    = []
    i = 9
    while i < len(raw):
        label = raw[i][0]
        if isinstance(label, str) and label.strip():
            label_clean = label.strip()
            if label_clean.lower() == 'sigma':
                break
            answers.append(label_clean)
            pct_row  = raw[i + 1] if i + 1 < len(raw) else []
            row_vals = []
            for j in brand_indices:
                v = pct_row[j] if j < len(pct_row) else None
                if v is None or (isinstance(v, float) and math.isnan(v)):
                    row_vals.append(None)
                elif isinstance(v, str) and v.strip() in ('-', '', '\xa0'):
                    row_vals.append(None)
                else:
                    row_vals.append(v)
            data.append(row_vals)
        i += 3

    return {
        'question_wording': question_wording,
        'brands':           brands,
        'base_values':      base_values,
        'answers':          answers,
        'data':             data,
    }


def parse_fmt5_sheet(sheet_df, desired_groups=None, weighted_data=False, weighted_base=None):
    if weighted_base is None:
        weighted_base = weighted_data
    """
    Parses the variant banner format where:
      Row 0  : project ID
      Row 1  : weight note
      Row 2  : banner/sample label
      Row 3  : question wording  (col 0)
      Row 4  : group names       (col 1+ — e.g. 'Total', blank...)
      Row 5  : category names    (col 1+ — e.g. 'Total','Gen Z','Male','Female'...)
      Row 6  : letter codes
      Row 7  : blank
      Row 8  : Unweighted Base + counts (same row)
      Row 9  : blank
      Row 10 : Base: Total Answering + weighted counts (same row)
      Row 11 : blank
      Row 12+: answer label + count (same row), % on next row, stat sig after
               every 3 rows, ends at 'Sigma'
    """
    raw = sheet_df.values.tolist()

    question_wording = str(raw[3][0]).strip() if (len(raw) > 3 and raw[3][0]) else ''

    # Categories are in row 5 (col 1+)
    cat_row = raw[5] if len(raw) > 5 else []
    all_cols = []
    for j in range(1, len(cat_row)):
        g = cat_row[j]
        if isinstance(g, str) and g.strip() and g.strip() != '\xa0':
            all_cols.append((j, g.strip(), ''))

    if desired_groups is not None:
        desired_lower = {d.lower() for d in desired_groups}
        selected_cols = [(j, g, s) for (j, g, s) in all_cols if g.lower() in desired_lower]
    else:
        selected_cols = all_cols  # include all by default

    col_indices = [j for (j, g, s) in selected_cols]
    col_labels  = [(g, s) for (j, g, s) in selected_cols]

    # Base values:
    # Row 8  = Unweighted Base
    # Row 10 = Base: Total Answering (weighted)
    # Use weighted_base to control which N= shows in headers
    base_row_idx = 10 if weighted_base else 8
    base_row     = raw[base_row_idx] if len(raw) > base_row_idx else []
    base_values  = [base_row[j] if j < len(base_row) else None for j in col_indices]

    # Answers start at row 12, every 3 rows
    answers  = []
    data     = []
    sig_data = []

    # Column letter codes at row 6
    letter_row  = raw[6] if len(raw) > 6 else []
    col_letters = [letter_row[j] if j < len(letter_row) else None for j in col_indices]

    i = 12
    while i < len(raw):
        label = raw[i][0]
        if isinstance(label, str) and label.strip():
            label_clean = label.strip()
            if label_clean.lower() == 'sigma':
                break
            answers.append(label_clean)
            pct_row = raw[i + 1] if i + 1 < len(raw) else []
            sig_row = raw[i + 2] if i + 2 < len(raw) else []
            row_vals = []
            sig_vals = []
            for j in col_indices:
                v = pct_row[j] if j < len(pct_row) else None
                if v is None or (isinstance(v, float) and math.isnan(v)):
                    row_vals.append(None)
                elif isinstance(v, str) and v.strip() in ('-', '', '\xa0'):
                    row_vals.append(None)
                else:
                    row_vals.append(v)
                sv = sig_row[j] if j < len(sig_row) else None
                sig_vals.append(sv if isinstance(sv, str) else None)
            data.append(row_vals)
            sig_data.append(sig_vals)
        i += 3

    return {
        'question_wording': question_wording,
        'columns':          col_labels,
        'base_values':      base_values,
        'answers':          answers,
        'data':             data,
        'sig_data':         sig_data,
        'col_letters':      col_letters,
    }


def parse_fmt6_sheet(sheet_df, desired_groups=None, weighted_data=False, weighted_base=None):
    """
    Like fmt5 but with an extra descriptor row at position 2.
      Row 0  : project ID
      Row 1  : weight note
      Row 2  : descriptor (e.g. 'Total sample', short label — NOT the question)
      Row 3  : question wording (col 0)
      Row 4  : group names (col 1+)
      Row 5  : category names (col 1+)
      Row 6  : letter codes
      Row 7  : blank
      Row 8  : Unweighted Base label (no counts on same row for blank template)
               OR Unweighted Base + counts (populated file)
      Row 9  : blank
      Row 10 : Base: Total Answering label / counts
      Row 11 : blank
      Row 12+: answer data, every 3 rows
    """
    if weighted_base is None:
        weighted_base = weighted_data

    raw = sheet_df.values.tolist()

    question_wording = str(raw[3][0]).strip() if (len(raw) > 3 and raw[3][0]) else ''

    # Categories at row 5
    cat_row  = raw[5] if len(raw) > 5 else []
    all_cols = []
    for j in range(1, len(cat_row)):
        g = cat_row[j]
        if isinstance(g, str) and g.strip() and g.strip() != '\xa0':
            all_cols.append((j, g.strip(), ''))

    if desired_groups is not None:
        desired_lower = {d.lower() for d in desired_groups}
        selected_cols = [(j, g, s) for (j, g, s) in all_cols if g.lower() in desired_lower]
    else:
        selected_cols = all_cols

    col_indices = [j for (j, g, s) in selected_cols]
    col_labels  = [(g, s) for (j, g, s) in selected_cols]

    # Base values: row 8 = unweighted, row 10 = weighted
    base_row_idx = 10 if weighted_base else 8
    base_row     = raw[base_row_idx] if len(raw) > base_row_idx else []
    base_values  = [base_row[j] if j < len(base_row) else None for j in col_indices]

    # Letter codes at row 6
    letter_row  = raw[6] if len(raw) > 6 else []
    col_letters = [letter_row[j] if j < len(letter_row) else None for j in col_indices]

    # Answers start at row 12, every 3 rows
    answers  = []
    data     = []
    sig_data = []
    i = 12
    while i < len(raw):
        label = raw[i][0]
        if isinstance(label, str) and label.strip():
            label_clean = label.strip()
            if label_clean.lower() == 'sigma':
                break
            answers.append(label_clean)
            pct_row = raw[i + 1] if i + 1 < len(raw) else []
            sig_row = raw[i + 2] if i + 2 < len(raw) else []
            row_vals = []
            sig_vals = []
            for j in col_indices:
                v = pct_row[j] if j < len(pct_row) else None
                if v is None or (isinstance(v, float) and math.isnan(v)):
                    row_vals.append(None)
                elif isinstance(v, str) and v.strip() in ('-', '', '\xa0'):
                    row_vals.append(None)
                else:
                    row_vals.append(v)
                sv = sig_row[j] if j < len(sig_row) else None
                sig_vals.append(sv if isinstance(sv, str) else None)
            data.append(row_vals)
            sig_data.append(sig_vals)
        i += 3

    return {
        'question_wording': question_wording,
        'columns':          col_labels,
        'base_values':      base_values,
        'answers':          answers,
        'data':             data,
        'sig_data':         sig_data,
        'col_letters':      col_letters,
    }


def parse_fmt7_sheet(sheet_df, desired_groups=None, weighted_data=False, weighted_base=None):
    """
    fmt7: weighted banner with extra group header row.
    Row 2: question wording
    Row 3: group names (Gender, Generation...) — may have blanks
    Row 4: category names (Total, Male, Female...) — col 1 = Total
    Row 5: letter codes
    Row 7: Unweighted Base + counts
    Row 9: Base: Total Answering + weighted counts
    Row 10+: data every 3 rows
    """
    if weighted_base is None:
        weighted_base = weighted_data

    raw = sheet_df.values.tolist()
    question_wording = str(raw[2][0]).strip() if (len(raw) > 2 and raw[2][0]) else ""

    # Categories at row 4
    cat_row = raw[4] if len(raw) > 4 else []
    all_cols = []
    for j in range(1, len(cat_row)):
        g = cat_row[j]
        if isinstance(g, str) and g.strip() and g.strip() != "\xa0":
            all_cols.append((j, g.strip(), ""))

    if desired_groups is not None:
        desired_lower = {d.lower() for d in desired_groups}
        selected_cols = [(j, g, s) for (j, g, s) in all_cols if g.lower() in desired_lower]
    else:
        selected_cols = [(j, g, s) for (j, g, s) in all_cols if True]

    col_indices = [j for (j, g, s) in selected_cols]
    col_labels  = [(g, s) for (j, g, s) in selected_cols]

    # Base rows: 7 = unweighted, 9 = weighted
    base_offset   = 2 if weighted_base else 0
    base_row_idx  = 7
    base_data_row = raw[base_row_idx + base_offset] if base_row_idx + base_offset < len(raw) else []
    base_values   = [base_data_row[j] if j < len(base_data_row) else None for j in col_indices]

    # Letter codes at row 5
    letter_row  = raw[5] if len(raw) > 5 else []
    col_letters = [letter_row[j] if j < len(letter_row) else None for j in col_indices]

    # Data starts at row 11 (7 + 4 for weighted, skipping both base rows)
    base_rows_used = 4 if weighted_data else 2
    data_start     = base_row_idx + base_rows_used

    answers  = []
    data     = []
    sig_data = []

    i = data_start
    while i < len(raw):
        label = raw[i][0]
        if isinstance(label, str) and label.strip():
            clean = label.strip()
            # Skip base-like rows that might appear
            if any(clean.lower().startswith(b) for b in ("base", "unweighted", "sigma")):
                i += 1
                continue
            if clean.lower() == "sigma":
                break
            answers.append(clean)
            pct_row = raw[i + 1] if i + 1 < len(raw) else []
            sig_row = raw[i + 2] if i + 2 < len(raw) else []
            row_vals = []
            sig_vals = []
            for j in col_indices:
                v = pct_row[j] if j < len(pct_row) else None
                if v is None or (isinstance(v, float) and math.isnan(v)):
                    row_vals.append(None)
                elif isinstance(v, str) and v.strip() in ("-", "", "\xa0"):
                    row_vals.append(None)
                else:
                    row_vals.append(v)
                sv = sig_row[j] if j < len(sig_row) else None
                sig_vals.append(sv if isinstance(sv, str) else None)
            data.append(row_vals)
            sig_data.append(sig_vals)
        i += 3

    return {
        "question_wording": question_wording,
        "columns":          col_labels,
        "base_values":      base_values,
        "answers":          answers,
        "data":             data,
        "sig_data":         sig_data,
        "col_letters":      col_letters,
    }


import re as _re

_FMT6_TABLE_TYPES = ['Summary Grid', 'Summary - Mean', 'T2B - Summary',
                     'B2B - Summary', 'T3B - Summary', 'HIDDEN']

def classify_fmt6_sheet(wording):
    # Normalize: remove apostrophes before brackets, lowercase for matching
    w = wording.replace("'[", "[")

    # Check for table type keywords anywhere in the wording
    table_type_map = [
        ('summary_grid',  ['summary grid', 'grid - summary', 'grid summary']),
        ('summary_mean',  ['summary - mean', "summary - mean'"]),
        ('t2b',           ['t2b - summary', "t2b - 'summary", "t2b - summary'"]),
        ('b2b',           ['b2b - summary', "b2b - 'summary", "b2b - summary'"]),
        ('t3b',           ['t3b - summary', "t3b - 'summary"]),
        ('hidden',        ['hidden']),
    ]
    wl = w.lower()
    for key, patterns in table_type_map:
        for pat in patterns:
            if pat in wl:
                return key, None

    # Entity sheet: extract first bracketed name that is NOT a table type or statement
    # For compound wordings like [Google] - [Summary Grid], find entity first
    matches = _re.findall(r'\[([^\[\]]+)\]', w)
    for m in matches:
        ml = m.strip().lower()
        # Skip if this bracket contains a table type keyword
        is_table = any(pat in ml for patterns in [p for _, p in table_type_map] for pat in patterns)
        if not is_table and len(ml) < 60:
            return 'entity', m.strip()

    return 'standalone', None


def parse_fmt6_mean(sheet_df, desired_groups=None, use_weighted_base=False):
    raw = sheet_df.values.tolist()
    question_wording = str(raw[3][0]).strip() if (len(raw) > 3 and raw[3][0]) else ''
    cat_row  = raw[5] if len(raw) > 5 else []
    all_cols = []
    for j in range(1, len(cat_row)):
        g = cat_row[j]
        if isinstance(g, str) and g.strip() and g.strip() != '\xa0':
            all_cols.append((j, g.strip()))
    if desired_groups:
        desired_lower = {d.lower() for d in desired_groups}
        selected = [(j, g) for j, g in all_cols if g.lower() in desired_lower]
    else:
        selected = all_cols
    col_indices = [j for j, g in selected]
    col_labels  = [g for j, g in selected]
    base_row  = raw[8] if len(raw) > 8 else []
    base_vals = [base_row[j] if j < len(base_row) else None for j in col_indices]
    answers, data = [], []
    SKIP = {'sigma', '- column means:', '- column proportions:', 'table of contents'}
    i = 10
    while i < len(raw):
        label = raw[i][0]
        if isinstance(label, str) and label.strip():
            clean = label.strip()
            if clean.lower() in SKIP:
                break
            answers.append(clean)
            row_vals = []
            for j in col_indices:
                v = raw[i][j] if j < len(raw[i]) else None
                if v is None or (isinstance(v, float) and math.isnan(v)):
                    row_vals.append(None)
                elif isinstance(v, str) and v.strip() in ('-', '', '\xa0'):
                    row_vals.append(None)
                else:
                    row_vals.append(v)
            data.append(row_vals)
        i += 2
    return {
        'question_wording': question_wording,
        'columns':          [(g, '') for g in col_labels],
        'col_labels':       col_labels,
        'base_values':      base_vals,
        'answers':          answers,
        'data':             data,
        'sig_data':         [[] for _ in answers],
        'col_letters':      [],
        'is_mean':          True,
    }


def build_fmt6_entity_merge(entity_sheets, desired_groups, weighted_data, use_weighted_base):
    TARGET = {'top 2 box', 'top 2 box (net)', 'bottom 2 box', 'bottom 2 box (net)'}
    col_labels, base_values = [], []
    ref_answers, merged_data = None, None
    for sheet_idx, entity_name, sheet_df in entity_sheets:
        parsed = parse_fmt6_sheet(sheet_df, desired_groups, weighted_data, use_weighted_base)
        if not parsed['answers']:
            continue
        col_list = [g for (g, s) in parsed['columns']]
        if ref_answers is None:
            ref_answers = parsed['answers']
            merged_data = [[] for _ in ref_answers]
        ans_lookup = {a.strip().lower(): i for i, a in enumerate(parsed['answers'])}
        for ci, col_label in enumerate(col_list):
            col_labels.append(f"{entity_name}\n{col_label}")
            base_values.append(parsed['base_values'][ci] if ci < len(parsed['base_values']) else None)
            for ai, ref_ans in enumerate(ref_answers):
                idx = ans_lookup.get(ref_ans.strip().lower())
                val = parsed['data'][idx][ci] if (idx is not None and idx < len(parsed['data']) and ci < len(parsed['data'][idx])) else None
                merged_data[ai].append(val)
    if not ref_answers:
        return None
    filtered_a, filtered_d = [], []
    for ai, ans in enumerate(ref_answers):
        if any(kw in ans.strip().lower() for kw in TARGET):
            filtered_a.append(ans.strip())
            filtered_d.append(merged_data[ai])
    return (col_labels, base_values, filtered_a, filtered_d) if filtered_a else None


def _hex_to_rgb(hex_color):
    h = hex_color.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def _interpolate_color(start_hex, end_hex, t):
    """Interpolate between two hex colors. t=0 → start, t=1 → end."""
    r1, g1, b1 = _hex_to_rgb(start_hex)
    r2, g2, b2 = _hex_to_rgb(end_hex)
    r = int(r1 + (r2 - r1) * t)
    g = int(g1 + (g2 - g1) * t)
    b = int(b1 + (b2 - b1) * t)
    return f"{r:02X}{g:02X}{b:02X}"

HEATMAP_SCHEMES = {
    "Blue scale":          ("#EFF6FF", "#1D4ED8"),
    "Green scale":         ("#F0FDF4", "#166534"),
    "Red-Green diverging": ("#DC2626", "#16A34A"),
    "None":                None,
}

def get_heatmap_color(value, min_val, max_val, scheme_name,
                      custom_start=None, custom_end=None):
    """Return hex color string (no #) for a value within range."""
    if scheme_name == "None" or min_val is None or max_val is None:
        return None
    if max_val == min_val:
        t = 0.5
    else:
        t = (value - min_val) / (max_val - min_val)
        t = max(0.0, min(1.0, t))

    if scheme_name == "Custom" and custom_start and custom_end:
        start, end = custom_start, custom_end
    else:
        pair = HEATMAP_SCHEMES.get(scheme_name)
        if not pair:
            return None
        start, end = pair

    return _interpolate_color(start, end, t)


def build_sig_flags(sig_data, col_letters, total_col_idx=0):
    """
    Build a 2D list of sig flags matching data shape.
    Flag = '▲' if this col's letter appears in sig_data of another col (higher than total)
         = '▼' if total col's sig_data contains this col's letter (lower than total)
         = ''  otherwise

    sig_data:     list of rows, each row = list of sig strings per col
    col_letters:  list of letter codes per col (e.g. ['A','B','C'...])
    total_col_idx: index of the Total column (default 0)
    """
    if not sig_data or not col_letters:
        return [['' for _ in row] for row in sig_data]

    total_letter = col_letters[total_col_idx] if total_col_idx < len(col_letters) else 'A'
    flags = []

    for row_sig in sig_data:
        row_flags = []
        # Get the total col's sig string for this row
        total_sig_str = row_sig[total_col_idx] if total_col_idx < len(row_sig) else ''
        total_sig_str = str(total_sig_str) if total_sig_str else ''

        for ci, sig_str in enumerate(row_sig):
            if ci == total_col_idx:
                row_flags.append('')
                continue
            sig_str = str(sig_str) if sig_str else ''
            this_letter = col_letters[ci] if ci < len(col_letters) else ''

            if total_letter and total_letter.upper() in sig_str.upper():
                # This col is sig higher than Total
                row_flags.append(' ▲')
            elif this_letter and this_letter.upper() in total_sig_str.upper():
                # Total is sig higher than this col → this col is sig lower
                row_flags.append(' ▼')
            else:
                row_flags.append('')
        flags.append(row_flags)
    return flags


# ── Word output ──────────────────────────────────────────────

def _get_word_template(portrait_landscape=False):
    import os
    template_file = 'template_portrait.docx' if portrait_landscape else 'template_landscape.docx'
    if os.path.exists(template_file):
        return Document(template_file)
    # Fallback to blank document if template not found
    return Document()


def write_table_to_doc(doc, question_wording, col_labels, base_values,
                       answers, data, multiple, is_first,
                       sig_flags=None, heatmap_scheme=None,
                       heatmap_custom_start=None, heatmap_custom_end=None):
    """
    sig_flags: 2D list [row][col] of '' / ' ▲' / ' ▼'
    heatmap_scheme: name from HEATMAP_SCHEMES or 'Custom'
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _set_cell_shading(cell, hex_color):
        """Apply background shading to a Word table cell."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

    if not is_first:
        doc.add_paragraph()
    q_para       = doc.add_paragraph()
    q_para.style = doc.styles['Normal']
    q_para.text  = question_wording

    n_cols = len(col_labels)
    table  = doc.add_table(rows=1, cols=n_cols + 1)
    table.style     = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    hdr[0].text = ''
    for i in range(n_cols):
        base_str = prettyPrint(base_values[i]) if i < len(base_values) else ''
        hdr[i + 1].text = ''
        hdr[i + 1].paragraphs[0].add_run(
            f"{col_labels[i]}\n(N={base_str})"
        ).bold = True
        hdr[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Collect all numeric values for heatmap range
    all_vals = []
    if heatmap_scheme and heatmap_scheme != "None":
        for row in data:
            for v in row:
                if v is not None and not isinstance(v, str):
                    try:
                        all_vals.append(float(v) * multiple)
                    except Exception:
                        pass
    min_val = min(all_vals) if all_vals else None
    max_val = max(all_vals) if all_vals else None

    for r_idx, answer in enumerate(answers):
        if str(answer).strip().lower() == 'sigma':
            continue
        row_cells = table.add_row().cells
        row_cells[0].text = str(answer)
        row_vals  = data[r_idx]     if r_idx < len(data)      else []
        row_flags = sig_flags[r_idx] if sig_flags and r_idx < len(sig_flags) else []

        for c_idx in range(n_cols):
            val   = row_vals[c_idx]  if c_idx < len(row_vals)  else None
            flag  = row_flags[c_idx] if c_idx < len(row_flags) else ''
            cell  = row_cells[c_idx + 1]
            cell.text = ''

            if val is None or (isinstance(val, float) and math.isnan(val)):
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif isinstance(val, str):
                cell.paragraphs[0].add_run(val + flag)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                pct_val = normal_round(round(val * multiple, 3))
                cell.paragraphs[0].add_run(f"{pct_val}%{flag}")
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Heatmap shading
                if heatmap_scheme and heatmap_scheme != "None" and min_val is not None:
                    hex_col = get_heatmap_color(
                        float(val) * multiple, min_val, max_val,
                        heatmap_scheme, heatmap_custom_start, heatmap_custom_end
                    )
                    if hex_col:
                        _set_cell_shading(cell, hex_col)

    for j in range(n_cols + 1):
        for cell in table.columns[j].cells:
            cell.width = Inches(3) if j == 0 else Inches(1.75)
            if j > 0:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# ── Excel output ─────────────────────────────────────────────

HDR_FILL  = PatternFill("solid", fgColor="1F4E79")
HDR_FONT  = XLFont(bold=True, color="FFFFFF", name="Arial", size=10)
BODY_FONT = XLFont(name="Arial", size=10)
CTR       = Alignment(horizontal="center", vertical="center", wrap_text=True)
LEFT      = Alignment(horizontal="left",   vertical="center", wrap_text=True)
THIN      = Side(style="thin", color="AAAAAA")
BORDER    = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


def _xl_write_table(ws, start_row, question_wording, col_headers,
                    base_values, answers, data, multiple=100,
                    show_base_in_header=True, sig_flags=None,
                    heatmap_scheme=None, heatmap_custom_start=None,
                    heatmap_custom_end=None):
    ws.cell(row=start_row, column=1, value=question_wording)
    ws.cell(row=start_row, column=1).font = XLFont(bold=True, name="Arial", size=10)
    ws.merge_cells(start_row=start_row, start_column=1,
                   end_row=start_row, end_column=len(col_headers) + 1)
    start_row += 1

    ws.cell(row=start_row, column=1, value='').fill   = HDR_FILL
    ws.cell(row=start_row, column=1).border = BORDER
    for ci, label in enumerate(col_headers):
        n_str = ''
        if show_base_in_header and ci < len(base_values) and base_values[ci] is not None:
            n_str = f"\n(N={prettyPrint(base_values[ci])})"
        cell            = ws.cell(row=start_row, column=ci + 2, value=label + n_str)
        cell.font       = HDR_FONT
        cell.fill       = HDR_FILL
        cell.alignment  = CTR
        cell.border     = BORDER
    start_row += 1

    # Heatmap range
    all_vals = []
    if heatmap_scheme and heatmap_scheme != "None":
        for row in data:
            for v in row:
                if v is not None and not isinstance(v, str):
                    try:
                        all_vals.append(float(v) * multiple)
                    except Exception:
                        pass
    min_val = min(all_vals) if all_vals else None
    max_val = max(all_vals) if all_vals else None

    for ri, answer in enumerate(answers):
        if str(answer).strip().lower() == 'sigma':
            continue
        row_vals   = data[ri]       if ri < len(data)       else []
        row_flags  = sig_flags[ri]  if sig_flags and ri < len(sig_flags) else []
        label_cell = ws.cell(row=start_row, column=1, value=str(answer))
        label_cell.font      = BODY_FONT
        label_cell.alignment = LEFT
        label_cell.border    = BORDER
        for ci in range(len(col_headers)):
            val   = row_vals[ci]  if ci < len(row_vals)  else None
            flag  = row_flags[ci] if ci < len(row_flags) else ''
            cell  = ws.cell(row=start_row, column=ci + 2)
            cell.font      = BODY_FONT
            cell.alignment = CTR
            cell.border    = BORDER
            if val is None or (isinstance(val, float) and math.isnan(val)):
                cell.value = ''
            elif isinstance(val, str):
                cell.value = val + (flag or '')
            else:
                pct_val    = normal_round(round(val * multiple, 3))
                cell.value = f"{pct_val}%{flag or ''}"
                # Heatmap fill
                if heatmap_scheme and heatmap_scheme != "None" and min_val is not None:
                    hex_col = get_heatmap_color(
                        float(val) * multiple, min_val, max_val,
                        heatmap_scheme, heatmap_custom_start, heatmap_custom_end
                    )
                    if hex_col:
                        cell.fill = PatternFill("solid", fgColor=hex_col)
        start_row += 1

    ws.column_dimensions['A'].width = max(ws.column_dimensions['A'].width, 35)
    for ci in range(len(col_headers)):
        col_letter = get_column_letter(ci + 2)
        ws.column_dimensions[col_letter].width = max(
            ws.column_dimensions[col_letter].width, 14)

    return start_row + 1


def _prep_fmt2_fmt5(parsed, row_filter, show_sig_flags):
    """Extract col_labels, filtered answers/data/sig, and compute flags."""
    col_labels  = [g for (g, s) in parsed['columns']]
    all_answers = parsed["answers"]
    all_data    = parsed["data"]
    all_sig     = parsed.get("sig_data", [[] for _ in all_answers])

    if row_filter and row_filter != 'all':
        custom = {r.strip().lower() for r in row_filter}
        filtered = [(a, d, s) for a, d, s in zip(all_answers, all_data, all_sig)
                    if a.strip().lower() in custom]
        answers  = [f[0] for f in filtered]
        data     = [f[1] for f in filtered]
        sig_data = [f[2] for f in filtered]
    else:
        answers  = all_answers
        data     = all_data
        sig_data = all_sig

    flags = None
    if show_sig_flags and sig_data:
        flags = build_sig_flags(sig_data, parsed.get('col_letters', []))

    return col_labels, answers, data, flags



def parse_fmt7_sheet(sheet_df, desired_groups=None, weighted_data=False, weighted_base=None):
    """fmt7: question at row2, group header at row3, categories at row4, weighted bases at rows 7+9."""
    if weighted_base is None:
        weighted_base = weighted_data
    raw = sheet_df.values.tolist()
    question_wording = str(raw[2][0]).strip() if (len(raw) > 2 and raw[2][0]) else ""
    cat_row = raw[4] if len(raw) > 4 else []
    all_cols = []
    for j in range(1, len(cat_row)):
        g = cat_row[j]
        if isinstance(g, str) and g.strip() and g.strip() != "\xa0":
            all_cols.append((j, g.strip(), ""))
    if desired_groups is not None:
        desired_lower = {d.lower() for d in desired_groups}
        selected_cols = [(j, g, s) for (j, g, s) in all_cols if g.lower() in desired_lower]
    else:
        selected_cols = all_cols
    col_indices = [j for (j, g, s) in selected_cols]
    col_labels  = [(g, s) for (j, g, s) in selected_cols]
    base_offset   = 2 if weighted_base else 0
    base_row_idx  = 7
    base_data_row = raw[base_row_idx + base_offset] if base_row_idx + base_offset < len(raw) else []
    base_values   = [base_data_row[j] if j < len(base_data_row) else None for j in col_indices]
    letter_row  = raw[5] if len(raw) > 5 else []
    col_letters = [letter_row[j] if j < len(letter_row) else None for j in col_indices]
    base_rows_used = 4 if weighted_data else 2
    data_start     = base_row_idx + base_rows_used
    answers, data, sig_data = [], [], []
    i = data_start
    while i < len(raw):
        label = raw[i][0]
        if isinstance(label, str) and label.strip():
            clean = label.strip()
            if any(clean.lower().startswith(b) for b in ("base", "unweighted", "sigma")):
                i += 1
                continue
            answers.append(clean)
            pct_row = raw[i + 1] if i + 1 < len(raw) else []
            sig_row = raw[i + 2] if i + 2 < len(raw) else []
            row_vals, sig_vals = [], []
            for j in col_indices:
                v = pct_row[j] if j < len(pct_row) else None
                if v is None or (isinstance(v, float) and math.isnan(v)):
                    row_vals.append(None)
                elif isinstance(v, str) and v.strip() in ("-", "", "\xa0"):
                    row_vals.append(None)
                else:
                    row_vals.append(v)
                sv = sig_row[j] if j < len(sig_row) else None
                sig_vals.append(sv if isinstance(sv, str) else None)
            data.append(row_vals)
            sig_data.append(sig_vals)
        i += 3
    return {"question_wording": question_wording, "columns": col_labels,
            "base_values": base_values, "answers": answers, "data": data,
            "sig_data": sig_data, "col_letters": col_letters}


def _generate_fmt6_outputs(
    xl, file_bytes, sheet_table_configs, desired_groups,
    output_format, excel_mode, portrait_landscape,
    weighted_data, use_weighted_base,
    out_merged, out_t2b, out_b2b, out_grid, out_mean, out_standalone,
    progress_callback,
):
    """Handles fmt6 files with entity grouping and merged tables."""
    skipped, errors = [], []
    groups, order = {}, []
    selected_indices = set()
    for cfg in sheet_table_configs:
        si = cfg[0] if not isinstance(cfg[0], str) else xl.sheet_names.index(cfg[0])
        selected_indices.add(si)

    for i in range(len(xl.sheet_names)):
        if i not in selected_indices:
            continue
        try:
            df  = xl.parse(i, header=None, na_values=[""])
            raw = df.values.tolist()
            wording = ""
            for row_idx in [3, 2]:
                if len(raw) > row_idx and raw[row_idx] and raw[row_idx][0]:
                    val = str(raw[row_idx][0]).strip()
                    if len(val) > 5 and "sample" not in val.lower() and "weight" not in val.lower():
                        wording = val
                        break
            if not wording:
                continue
            sheet_type, entity = classify_fmt6_sheet(wording)
            if sheet_type == "hidden":
                continue
            m = _re.match(r"^([A-Za-z0-9_]+)\.", wording.strip())
            prefix = m.group(1) if m else wording[:20]
            if prefix not in groups:
                groups[prefix] = {
                    "wording_base": _re.sub(r"\s*\[.*?\]\s*-?\s*", " ", wording).strip(),
                    "entity": [], "t2b": [], "b2b": [], "t3b": [],
                    "summary_grid": [], "summary_mean": [], "standalone": [],
                }
                order.append(prefix)
            entry = (i, entity, df)
            if sheet_type == "entity":
                groups[prefix]["entity"].append(entry)
            elif "t2b" in sheet_type:
                groups[prefix]["t2b"].append(entry)
            elif "b2b" in sheet_type:
                groups[prefix]["b2b"].append(entry)
            elif sheet_type == "summary_grid":
                groups[prefix]["summary_grid"].append(entry)
            elif "mean" in sheet_type:
                groups[prefix]["summary_mean"].append(entry)
            else:
                groups[prefix]["standalone"].append(entry)
        except Exception as e:
            errors.append((xl.sheet_names[i], str(e)))

    word_doc = None
    if output_format in ("word", "both"):
        try:
            word_doc = _get_word_template(portrait_landscape)
        except Exception:
            word_doc = Document()
        word_doc.styles["Normal"].font.name = "Arial"
        word_doc.styles["Normal"].font.size = Pt(10)

    xl_wb, xl_sheet_rows, xl_sheet_ctr = None, {}, [0]
    if output_format in ("excel", "both"):
        xl_wb = openpyxl.Workbook()
        xl_wb.remove(xl_wb.active)

    def _xl_get(label):
        name = label[:31]
        if excel_mode == "per_question":
            if name not in xl_wb.sheetnames:
                xl_wb.create_sheet(title=name)
                xl_sheet_rows[name] = 1
            return xl_wb[name], xl_sheet_rows.get(name, 1)
        else:
            xl_sheet_ctr[0] += 1
            sname = f"{xl_sheet_ctr[0]:03d}_{name}"[:31]
            return xl_wb.create_sheet(title=sname), 1

    def _xl_done(ws, label, nr):
        if excel_mode == "per_question":
            xl_sheet_rows[label[:31]] = nr

    first_word = [True]

    def _write(q_wording, col_labels, base_values, answers, data, xl_key, is_mean=False):
        multiple = 1 if is_mean else 100
        if word_doc is not None:
            if not first_word[0]:
                word_doc.add_paragraph()
            qp = word_doc.add_paragraph()
            qp.style = word_doc.styles["Normal"]
            qp.text  = q_wording
            n_cols = len(col_labels)
            tbl = word_doc.add_table(rows=1, cols=n_cols + 1)
            tbl.style     = "Table Grid"
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = tbl.rows[0].cells
            hdr[0].text = ""
            for ci in range(n_cols):
                base_str = prettyPrint(base_values[ci]) if ci < len(base_values) else ""
                hdr[ci+1].text = ""
                run = hdr[ci+1].paragraphs[0].add_run(f"{col_labels[ci]}\n(N={base_str})")
                run.bold = True
                hdr[ci+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for ri, ans in enumerate(answers):
                rc = tbl.add_row().cells
                rc[0].text = str(ans)
                row_vals = data[ri] if ri < len(data) else []
                for ci in range(n_cols):
                    val  = row_vals[ci] if ci < len(row_vals) else None
                    cell = rc[ci + 1]
                    cell.text = ""
                    if val is None or (isinstance(val, float) and math.isnan(val)):
                        pass
                    elif isinstance(val, str):
                        cell.paragraphs[0].add_run(val)
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif is_mean:
                        cell.paragraphs[0].add_run(str(round(val, 2)))
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        cell.paragraphs[0].add_run(f"{normal_round(round(val * 100, 3))}%")
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for j in range(n_cols + 1):
                for cell in tbl.columns[j].cells:
                    cell.width = Inches(2.5) if j == 0 else Inches(1.4)
                    if j > 0:
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            first_word[0] = False

        if xl_wb is not None:
            ws, row_num = _xl_get(xl_key)
            nr = _xl_write_table(ws, row_num, q_wording, col_labels,
                                 base_values, answers, data, multiple)
            _xl_done(ws, xl_key, nr)

    total_groups = max(len(order), 1)
    for gi, prefix in enumerate(order):
        if progress_callback:
            progress_callback(gi / total_groups, f"Processing {prefix}…")
        grp    = groups[prefix]
        xl_key = prefix[:31]

        if out_merged and grp["entity"]:
            result = build_fmt6_entity_merge(grp["entity"], desired_groups, weighted_data, use_weighted_base)
            if result:
                cls, bvs, ans, dat = result
                _write(grp["wording_base"] + " [Companies T2B/B2B]", cls, bvs, ans, dat, xl_key)

        type_cfg = [
            ("t2b",          out_t2b,  False, " [T2B Summary]"),
            ("b2b",          out_b2b,  False, " [B2B Summary]"),
            ("summary_grid", out_grid, False, " [Summary Grid]"),
            ("summary_mean", out_mean, True,  " [Mean]"),
        ]
        for key, enabled, use_mean, suffix in type_cfg:
            if not enabled:
                continue
            for si, entity, df in grp[key]:
                try:
                    if use_mean:
                        p = parse_fmt6_mean(df, desired_groups, use_weighted_base)
                    elif key == "summary_grid":
                        p = parse_fmt6_sheet(df, None, weighted_data, use_weighted_base)
                    else:
                        p = parse_fmt6_sheet(df, desired_groups, weighted_data, use_weighted_base)
                    if not p["answers"]:
                        continue
                    cls = p.get("col_labels") or [g for g, s in p["columns"]]
                    _write(p["question_wording"] + suffix, cls, p["base_values"],
                           p["answers"], p["data"], xl_key, is_mean=p.get("is_mean", False))
                except Exception as e:
                    errors.append((xl.sheet_names[si], str(e)))

        if out_standalone:
            for si, entity, df in grp["standalone"]:
                try:
                    p = parse_fmt6_sheet(df, desired_groups, weighted_data, use_weighted_base)
                    if not p["answers"]:
                        continue
                    cls = p.get("col_labels") or [g for g, s in p["columns"]]
                    _write(p["question_wording"], cls, p["base_values"],
                           p["answers"], p["data"], xl_key)
                except Exception as e:
                    errors.append((xl.sheet_names[si], str(e)))

    if progress_callback:
        progress_callback(1.0, "Finalizing…")

    word_bytes = excel_bytes = None
    if word_doc is not None and output_format in ("word", "both"):
        buf = io.BytesIO()
        word_doc.save(buf)
        word_bytes = buf.getvalue()
    if xl_wb is not None and output_format in ("excel", "both"):
        buf = io.BytesIO()
        xl_wb.save(buf)
        excel_bytes = buf.getvalue()

    return {"word_bytes": word_bytes, "excel_bytes": excel_bytes,
            "skipped": skipped, "errors": errors}


# ── Main generation function ─────────────────────────────────

def generate_outputs(
    file_bytes,
    sheet_table_configs,
    desired_groups,
    output_format,
    excel_mode,
    portrait_landscape,
    weighted_data,
    use_weighted_base=None,
    progress_callback=None,
    show_sig_flags=False,
    heatmap_scheme="None",
    heatmap_custom_start=None,
    heatmap_custom_end=None,
    fmt6_output_merged_entity=True,
    fmt6_output_t2b=True,
    fmt6_output_b2b=True,
    fmt6_output_grid=True,
    fmt6_output_mean=False,
    fmt6_output_standalone=True,
):
    if use_weighted_base is None:
        use_weighted_base = weighted_data

    xl    = pd.ExcelFile(io.BytesIO(file_bytes))
    total = max(len(sheet_table_configs), 1)

    # Detect fmt6-dominant file and route to dedicated handler
    fmt6_count = 0
    for cfg in sheet_table_configs[:10]:
        try:
            si  = cfg[0] if not isinstance(cfg[0], str) else xl.sheet_names.index(cfg[0])
            df  = xl.parse(si, header=None, na_values=[''])
            fmt = detect_format(df)
            if fmt == 'fmt6':
                fmt6_count += 1
        except Exception:
            pass

    if fmt6_count >= 3:
        return _generate_fmt6_outputs(
            xl, file_bytes, sheet_table_configs, desired_groups,
            output_format, excel_mode, portrait_landscape,
            weighted_data, use_weighted_base,
            fmt6_output_merged_entity, fmt6_output_t2b, fmt6_output_b2b,
            fmt6_output_grid, fmt6_output_mean, fmt6_output_standalone,
            progress_callback,
        )

    skipped = []
    errors  = []
    parsed_cache = {}

    # Word setup
    word_doc = None
    if output_format in ('word', 'both'):
        try:
            word_doc = _get_word_template(portrait_landscape)
        except Exception:
            word_doc = Document()
        style = word_doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)

    # Excel setup
    xl_wb         = None
    xl_sheet_ctr  = [0]
    xl_sheet_rows = {}

    if output_format in ('excel', 'both'):
        xl_wb = openpyxl.Workbook()
        xl_wb.remove(xl_wb.active)

    def _xl_get_sheet(label):
        name = label[:31]
        if excel_mode == 'per_question':
            if name not in xl_wb.sheetnames:
                xl_wb.create_sheet(title=name)
                xl_sheet_rows[name] = 1
            return xl_wb[name], xl_sheet_rows[name]
        else:
            xl_sheet_ctr[0] += 1
            sname = f"{xl_sheet_ctr[0]:03d}_{name}"[:31]
            return xl_wb.create_sheet(title=sname), 1

    def _xl_done(ws, label, next_row):
        if excel_mode == 'per_question':
            xl_sheet_rows[label[:31]] = next_row

    def write_xl(question_wording, col_headers, base_values, answers,
                 data, multiple, sheet_label, show_base=True):
        ws, row = _xl_get_sheet(sheet_label)
        nr = _xl_write_table(ws, row, question_wording, col_headers,
                             base_values, answers, data, multiple, show_base)
        _xl_done(ws, sheet_label, nr)

    def apply_row_filter(answers, data, row_filter):
        if row_filter == 'all' or not row_filter:
            return answers, data
        custom = {r.strip().lower() for r in row_filter}
        fa, fd = [], []
        for a, d in zip(answers, data):
            if a.strip().lower() in custom:
                fa.append(a)
                fd.append(d)
        return fa, fd

    first_word = True

    for idx, config_tuple in enumerate(sheet_table_configs):
        # Support both old (sheet_idx, row_filter) and new (sheet_idx, row_filter, prefix, label)
        if len(config_tuple) == 4:
            sheet_idx, row_filter, group_prefix, table_label = config_tuple
        else:
            sheet_idx, row_filter = config_tuple
            group_prefix = None
            table_label  = None

        if progress_callback:
            progress_callback(idx / total, f"Processing sheet {sheet_idx}…")

        try:
            # sheet_idx can be an integer (position) or string (sheet name)
            if isinstance(sheet_idx, str):
                # Find position by name
                sheet_name = sheet_idx
                sheet_idx  = xl.sheet_names.index(sheet_idx)
            else:
                sheet_name = xl.sheet_names[sheet_idx]
            sheet_label = sheet_name[:31]

            # Parse once per sheet, cache result
            if sheet_idx not in parsed_cache:
                df  = xl.parse(sheet_idx, header=None, na_values=[''])
                fmt = detect_format(df)
                parsed_cache[sheet_idx] = (df, fmt)
            sheet_df, fmt = parsed_cache[sheet_idx]

            # Build xl_sheet_key from question wording for readable sheet names
            # Use group_prefix if provided, otherwise extract from question wording
            if group_prefix:
                xl_sheet_key = group_prefix[:31]
            else:
                # Try to get question wording from raw data
                try:
                    raw_preview = sheet_df.values.tolist()
                    wording_row = 3 if fmt == 'fmt5' else 2
                    wording = str(raw_preview[wording_row][0]).strip() if (
                        len(raw_preview) > wording_row and raw_preview[wording_row][0]
                    ) else sheet_name
                    # Clean up: remove special chars that Excel doesn't allow in sheet names
                    import re as _re
                    clean = _re.sub(r'[\\/*?\[\]:]', '', wording)[:31].strip()
                    xl_sheet_key = clean if clean else sheet_name[:31]
                except Exception:
                    xl_sheet_key = sheet_name[:31]

            # ── fmt2 ─────────────────────────────────────────
            if fmt == 'fmt2':
                parsed  = parse_fmt2_sheet(sheet_df, desired_groups, weighted_data, weighted_base=use_weighted_base)
                if not parsed['answers']:
                    skipped.append(sheet_label)
                    continue

                col_labels, answers, data, flags = _prep_fmt2_fmt5(
                    parsed, row_filter, show_sig_flags)
                if not answers:
                    skipped.append(sheet_label)
                    continue

                if word_doc is not None:
                    write_table_to_doc(
                        word_doc, parsed['question_wording'],
                        col_labels, parsed['base_values'],
                        answers, data, 100, first_word,
                        sig_flags=flags,
                        heatmap_scheme=heatmap_scheme,
                        heatmap_custom_start=heatmap_custom_start,
                        heatmap_custom_end=heatmap_custom_end,
                    )
                    first_word = False

                if xl_wb is not None:
                    ws, row_num = _xl_get_sheet(xl_sheet_key)
                    nr = _xl_write_table(
                        ws, row_num, parsed['question_wording'],
                        col_labels, parsed['base_values'],
                        answers, data, 100, True,
                        sig_flags=flags,
                        heatmap_scheme=heatmap_scheme,
                        heatmap_custom_start=heatmap_custom_start,
                        heatmap_custom_end=heatmap_custom_end,
                    )
                    _xl_done(ws, xl_sheet_key, nr)

            # ── fmt5 ─────────────────────────────────────────
            elif fmt == 'fmt5':
                parsed = parse_fmt5_sheet(sheet_df, desired_groups, weighted_data, weighted_base=use_weighted_base)
                if not parsed['answers']:
                    skipped.append(sheet_label)
                    continue

                col_labels, answers, data, flags = _prep_fmt2_fmt5(
                    parsed, row_filter, show_sig_flags)
                if not answers:
                    skipped.append(sheet_label)
                    continue

                if word_doc is not None:
                    write_table_to_doc(
                        word_doc, parsed['question_wording'],
                        col_labels, parsed['base_values'],
                        answers, data, 100, first_word,
                        sig_flags=flags,
                        heatmap_scheme=heatmap_scheme,
                        heatmap_custom_start=heatmap_custom_start,
                        heatmap_custom_end=heatmap_custom_end,
                    )
                    first_word = False

                if xl_wb is not None:
                    ws, row_num = _xl_get_sheet(xl_sheet_key)
                    nr = _xl_write_table(
                        ws, row_num, parsed['question_wording'],
                        col_labels, parsed['base_values'],
                        answers, data, 100, True,
                        sig_flags=flags,
                        heatmap_scheme=heatmap_scheme,
                        heatmap_custom_start=heatmap_custom_start,
                        heatmap_custom_end=heatmap_custom_end,
                    )
                    _xl_done(ws, xl_sheet_key, nr)

            # ── fmt6 ─────────────────────────────────────────
            elif fmt == 'fmt6':
                parsed = parse_fmt6_sheet(sheet_df, desired_groups, weighted_data, weighted_base=use_weighted_base)
                if not parsed['answers']:
                    skipped.append(sheet_label)
                    continue

                col_labels, answers, data, flags = _prep_fmt2_fmt5(
                    parsed, row_filter, show_sig_flags)
                if not answers:
                    skipped.append(sheet_label)
                    continue

                if word_doc is not None:
                    write_table_to_doc(
                        word_doc, parsed['question_wording'],
                        col_labels, parsed['base_values'],
                        answers, data, 100, first_word,
                        sig_flags=flags,
                        heatmap_scheme=heatmap_scheme,
                        heatmap_custom_start=heatmap_custom_start,
                        heatmap_custom_end=heatmap_custom_end,
                    )
                    first_word = False

                if xl_wb is not None:
                    ws, row_num = _xl_get_sheet(xl_sheet_key)
                    nr = _xl_write_table(
                        ws, row_num, parsed['question_wording'],
                        col_labels, parsed['base_values'],
                        answers, data, 100, True,
                        sig_flags=flags,
                        heatmap_scheme=heatmap_scheme,
                        heatmap_custom_start=heatmap_custom_start,
                        heatmap_custom_end=heatmap_custom_end,
                    )
                    _xl_done(ws, xl_sheet_key, nr)

            # ── fmt7 ─────────────────────────────────────────
            elif fmt == 'fmt7':
                parsed = parse_fmt7_sheet(sheet_df, desired_groups, weighted_data, weighted_base=use_weighted_base)
                if not parsed['answers']:
                    skipped.append(sheet_label)
                    continue

                col_labels, answers, data, flags = _prep_fmt2_fmt5(
                    parsed, row_filter, show_sig_flags)
                if not answers:
                    skipped.append(sheet_label)
                    continue

                if word_doc is not None:
                    write_table_to_doc(
                        word_doc, parsed['question_wording'],
                        col_labels, parsed['base_values'],
                        answers, data, 100, first_word,
                        sig_flags=flags,
                        heatmap_scheme=heatmap_scheme,
                        heatmap_custom_start=heatmap_custom_start,
                        heatmap_custom_end=heatmap_custom_end,
                    )
                    first_word = False

                if xl_wb is not None:
                    ws, row_num = _xl_get_sheet(xl_sheet_key)
                    nr = _xl_write_table(
                        ws, row_num, parsed['question_wording'],
                        col_labels, parsed['base_values'],
                        answers, data, 100, True,
                        sig_flags=flags,
                        heatmap_scheme=heatmap_scheme,
                        heatmap_custom_start=heatmap_custom_start,
                        heatmap_custom_end=heatmap_custom_end,
                    )
                    _xl_done(ws, xl_sheet_key, nr)


            # ── fmt7 ─────────────────────────────────────────
            elif fmt == 'fmt7':
                parsed = parse_fmt7_sheet(sheet_df, desired_groups, weighted_data, weighted_base=use_weighted_base)
                if not parsed['answers']:
                    skipped.append(sheet_label)
                    continue
                col_labels, answers, data, flags = _prep_fmt2_fmt5(parsed, row_filter, show_sig_flags)
                if not answers:
                    skipped.append(sheet_label)
                    continue
                if word_doc is not None:
                    write_table_to_doc(
                        word_doc, parsed['question_wording'],
                        col_labels, parsed['base_values'],
                        answers, data, 100, first_word,
                        sig_flags=flags,
                        heatmap_scheme=heatmap_scheme,
                        heatmap_custom_start=heatmap_custom_start,
                        heatmap_custom_end=heatmap_custom_end,
                    )
                    first_word = False
                if xl_wb is not None:
                    ws, row_num = _xl_get_sheet(xl_sheet_key)
                    nr = _xl_write_table(
                        ws, row_num, parsed['question_wording'],
                        col_labels, parsed['base_values'],
                        answers, data, 100, True,
                        sig_flags=flags,
                        heatmap_scheme=heatmap_scheme,
                        heatmap_custom_start=heatmap_custom_start,
                        heatmap_custom_end=heatmap_custom_end,
                    )
                    _xl_done(ws, xl_sheet_key, nr)

            elif fmt == 'fmt3':
                parsed     = parse_fmt3_sheet(sheet_df, desired_groups)
                if not parsed['answers']:
                    skipped.append(sheet_label)
                    continue

                col_labels    = [g for (g, s) in parsed['columns']]
                companies     = parsed['answers']
                company_bases = parsed['company_bases']
                pct_data      = parsed['data']
                n_cols        = len(col_labels)

                # Apply row filter to companies
                if row_filter not in ('all', None):
                    companies, pct_data, company_bases = zip(
                        *[(c, d, b) for c, d, b in zip(companies, pct_data, company_bases)
                          if row_filter == 'all' or c.strip().lower() in {r.lower() for r in (row_filter if isinstance(row_filter, list) else [])}]
                    ) if companies else ([], [], [])
                    companies     = list(companies)
                    pct_data      = list(pct_data)
                    company_bases = list(company_bases)

                if not companies:
                    skipped.append(sheet_label)
                    continue

                def _write_fmt3_word(show_n, is_first_arg, suffix=''):
                    if word_doc is None:
                        return
                    if not is_first_arg:
                        word_doc.add_paragraph()
                    qp       = word_doc.add_paragraph()
                    qp.style = word_doc.styles['Normal']
                    qp.text  = parsed['question_wording'] + suffix
                    tbl = word_doc.add_table(rows=1, cols=n_cols + 1)
                    tbl.style     = 'Table Grid'
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    hdr = tbl.rows[0].cells
                    hdr[0].text = ''
                    for ci in range(n_cols):
                        hdr[ci+1].text = ''
                        hdr[ci+1].paragraphs[0].add_run(col_labels[ci]).bold = True
                        hdr[ci+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for ri, company in enumerate(companies):
                        rc   = tbl.add_row().cells
                        rc[0].text = company
                        bases = company_bases[ri] if ri < len(company_bases) else []
                        pcts  = pct_data[ri]      if ri < len(pct_data)      else []
                        for ci in range(n_cols):
                            pct  = pcts[ci]  if ci < len(pcts)  else None
                            n    = bases[ci] if ci < len(bases) else None
                            cell = rc[ci + 1]
                            cell.text = ''
                            if pct is None or isinstance(pct, str):
                                cell.paragraphs[0].add_run('-')
                            else:
                                pct_str = f"{normal_round(round(pct * 100, 3))}%"
                                n_str   = f"\n(N={prettyPrint(n)})" if (show_n and n) else ''
                                cell.paragraphs[0].add_run(pct_str + n_str)
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for ci in range(n_cols + 1):
                        for cell in tbl.columns[ci].cells:
                            cell.width = Inches(3) if ci == 0 else Inches(1.75)

                _write_fmt3_word(show_n=False, is_first_arg=first_word)
                _write_fmt3_word(show_n=True,  is_first_arg=False, suffix=' (with base)')
                first_word = False

                if xl_wb is not None:
                    write_xl(parsed['question_wording'], col_labels,
                             [None]*n_cols, companies, pct_data,
                             100, xl_sheet_key, show_base=False)
                    # second table with N
                    data_with_n = []
                    for ri, company in enumerate(companies):
                        bases = company_bases[ri] if ri < len(company_bases) else []
                        pcts  = pct_data[ri]
                        row   = []
                        for ci in range(n_cols):
                            pct = pcts[ci]  if ci < len(pcts)  else None
                            n   = bases[ci] if ci < len(bases) else None
                            if pct is None or isinstance(pct, str):
                                row.append(None)
                            else:
                                pct_str = f"{normal_round(round(pct * 100, 3))}%"
                                n_str   = f" (N={prettyPrint(n)})" if n else ''
                                row.append(pct_str + n_str)
                        data_with_n.append(row)
                    write_xl(parsed['question_wording'] + ' (with base)',
                             col_labels, [None]*n_cols, companies,
                             data_with_n, 1, xl_sheet_key, show_base=False)

            # ── fmt4 ─────────────────────────────────────────
            elif fmt == 'fmt4':
                parsed  = parse_fmt4_sheet(sheet_df)
                if not parsed['answers']:
                    skipped.append(sheet_label)
                    continue

                answers, data = apply_row_filter(parsed["answers"], parsed["data"], row_filter)
                if not answers:
                    skipped.append(sheet_label)
                    continue

                brands      = parsed['brands']
                base_values = parsed['base_values']
                n_brands    = len(brands)

                if word_doc is not None:
                    if not first_word:
                        word_doc.add_paragraph()
                    qp       = word_doc.add_paragraph()
                    qp.style = word_doc.styles['Normal']
                    qp.text  = parsed['question_wording']
                    tbl = word_doc.add_table(rows=1, cols=n_brands + 1)
                    tbl.style     = 'Table Grid'
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    hdr = tbl.rows[0].cells
                    hdr[0].text = ''
                    for ci, brand in enumerate(brands):
                        n = base_values[ci] if ci < len(base_values) else None
                        hdr[ci+1].text = ''
                        hdr[ci+1].paragraphs[0].add_run(
                            f"{brand}\n(N={prettyPrint(n)})"
                        ).bold = True
                        hdr[ci+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for ri, answer in enumerate(answers):
                        rc = tbl.add_row().cells
                        rc[0].text = answer
                        row_vals = data[ri] if ri < len(data) else []
                        for ci in range(n_brands):
                            val  = row_vals[ci] if ci < len(row_vals) else None
                            cell = rc[ci + 1]
                            cell.text = ''
                            if val is None:
                                cell.paragraphs[0].add_run('-')
                            else:
                                cell.paragraphs[0].add_run(
                                    f"{normal_round(round(val * 100, 3))}%"
                                )
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for ci in range(n_brands + 1):
                        for cell in tbl.columns[ci].cells:
                            cell.width = Inches(2) if ci == 0 else Inches(1.25)
                    first_word = False

                if xl_wb is not None:
                    write_xl(parsed['question_wording'], brands,
                             base_values, answers, data, 100, xl_sheet_key)

            else:
                skipped.append(f"{sheet_label} (fmt1/grid — skipped)")

        except Exception as e:
            errors.append((sheet_name if 'sheet_name' in dir() else str(sheet_idx), str(e)))

    if progress_callback:
        progress_callback(1.0, "Finalizing…")

    # Serialize outputs
    word_bytes  = None
    excel_bytes = None

    if word_doc is not None and output_format in ('word', 'both'):
        buf = io.BytesIO()
        word_doc.save(buf)
        word_bytes = buf.getvalue()

    if xl_wb is not None and output_format in ('excel', 'both'):
        buf = io.BytesIO()
        xl_wb.save(buf)
        excel_bytes = buf.getvalue()

    return {
        'word_bytes':  word_bytes,
        'excel_bytes': excel_bytes,
        'skipped':     skipped,
        'errors':      errors,
    }

# ── Multi-source scanning ─────────────────────────────────────

def scan_multi_source(file_list):
    """
    Scan multiple banner files and return a unified question index.
    file_list: list of (file_bytes, banner_name) tuples.

    Returns:
      {
        'banners': [
            {'name': str, 'sheets': [meta...], 'all_cols': [...]}
        ],
        'matched':   [ {q_id, wording, banner_sheets: {banner_name: sheet_meta}} ],
        'unmatched': [ {banner_name, sheet_meta} ]  # questions only in one banner
      }
    """
    import re
    import difflib

    banners = []
    for file_bytes, banner_name in file_list:
        metas    = scan_file(file_bytes)
        all_cols = get_all_columns(metas)
        banners.append({'name': banner_name, 'sheets': metas,
                        'all_cols': all_cols, 'file_bytes': file_bytes})

    if not banners:
        return {'banners': [], 'matched': [], 'unmatched': []}

    def q_id(wording):
        """Extract question ID prefix e.g. 'GO2', 'A0', 'S1'."""
        m = re.match(r'^([A-Za-z]+\d+)', wording.strip())
        return m.group(1).upper() if m else None

    # Build index per banner: {q_id or wording_key → sheet_meta}
    banner_indices = []
    for b in banners:
        idx = {}
        for m in b['sheets']:
            if m['fmt'] == 'error':
                continue
            qid = q_id(m['question_wording'])
            key = qid if qid else m['question_wording'].strip()[:80]
            # If duplicate key within same banner, append index
            if key in idx:
                key = f"{key}_{m['index']}"
            idx[key] = m
        banner_indices.append(idx)

    # Match questions across all banners
    # Start from banner 0 as the reference
    ref_idx  = banner_indices[0]
    matched  = []
    unmatched = []

    for key, ref_meta in ref_idx.items():
        entry = {
            'q_id':    key,
            'wording': ref_meta['question_wording'],
            'banner_sheets': {banners[0]['name']: ref_meta},
        }
        for bi in range(1, len(banners)):
            other_idx = banner_indices[bi]
            if key in other_idx:
                entry['banner_sheets'][banners[bi]['name']] = other_idx[key]
            else:
                # Not found by ID — mark as None (user will map manually)
                entry['banner_sheets'][banners[bi]['name']] = None
        matched.append(entry)

    # Find questions only in non-reference banners
    for bi in range(1, len(banners)):
        other_idx = banner_indices[bi]
        for key, meta in other_idx.items():
            if key not in ref_idx:
                unmatched.append({'banner_name': banners[bi]['name'],
                                  'q_id': key, 'sheet_meta': meta})

    return {
        'banners':   banners,
        'matched':   matched,
        'unmatched': unmatched,
    }


def generate_merged_outputs(
    multi_source,          # result of scan_multi_source
    matched_overrides,     # dict: {q_id → {banner_name: sheet_index|None}}
                           # user-provided manual mappings for unmatched questions
    selected_cols_per_banner,  # dict: {banner_name: [col_name, ...]}
    output_format,
    excel_mode,
    portrait_landscape,
    weighted_data,
    row_filter,            # 'all' or list of row labels
    progress_callback=None,
):
    """
    Generates merged output for Scenario A (column merge, same wave).
    For each matched question, columns from all banners are merged side by side.
    """
    banners  = multi_source['banners']
    matched  = multi_source['matched']
    skipped  = []
    errors   = []

    # Apply manual overrides to matched list
    # Override format: {q_id: {banner_name: sheet_index or None}}
    effective_matched = []
    for entry in matched:
        eff = dict(entry)
        if entry['q_id'] in (matched_overrides or {}):
            overrides = matched_overrides[entry['q_id']]
            for bn, sheet_idx in overrides.items():
                if sheet_idx is None:
                    eff['banner_sheets'][bn] = None
                else:
                    # Find the sheet meta by index
                    for b in banners:
                        if b['name'] == bn:
                            for m in b['sheets']:
                                if m['index'] == sheet_idx:
                                    eff['banner_sheets'][bn] = m
        effective_matched.append(eff)

    # Add unmatched questions that were manually mapped
    for entry in multi_source.get('unmatched', []):
        qid = entry['q_id']
        if qid in (matched_overrides or {}):
            overrides = matched_overrides[qid]
            new_entry = {
                'q_id':    qid,
                'wording': entry['sheet_meta']['question_wording'],
                'banner_sheets': {},
            }
            for b in banners:
                new_entry['banner_sheets'][b['name']] = None
            new_entry['banner_sheets'][entry['banner_name']] = entry['sheet_meta']
            for bn, sheet_idx in overrides.items():
                if sheet_idx is not None:
                    for b in banners:
                        if b['name'] == bn:
                            for m in b['sheets']:
                                if m['index'] == sheet_idx:
                                    new_entry['banner_sheets'][bn] = m
            effective_matched.append(new_entry)

    # Word + Excel setup
    word_doc = None
    if output_format in ('word', 'both'):
        try:
            word_doc = _get_word_template(portrait_landscape)
        except Exception:
            word_doc = Document()
        style = word_doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)

    xl_wb         = None
    xl_sheet_ctr  = [0]
    xl_sheet_rows = {}

    if output_format in ('excel', 'both'):
        xl_wb = openpyxl.Workbook()
        xl_wb.remove(xl_wb.active)

    def _xl_get_sheet(label):
        name = label[:31]
        if excel_mode == 'per_question':
            if name not in xl_wb.sheetnames:
                xl_wb.create_sheet(title=name)
                xl_sheet_rows[name] = 1
            return xl_wb[name], xl_sheet_rows.get(name, 1)
        else:
            xl_sheet_ctr[0] += 1
            sname = f"{xl_sheet_ctr[0]:03d}_{name}"[:31]
            return xl_wb.create_sheet(title=sname), 1

    def _xl_done(ws, label, next_row):
        if excel_mode == 'per_question':
            xl_sheet_rows[label[:31]] = next_row

    def _apply_row_filter(answers, data):
        if row_filter == 'all' or not row_filter:
            return answers, data
        custom = {r.strip().lower() for r in row_filter}
        fa, fd = [], []
        for a, d in zip(answers, data):
            if a.strip().lower() in custom:
                fa.append(a)
                fd.append(d)
        return fa, fd

    total     = max(len(effective_matched), 1)
    first_doc = True

    # File bytes lookup
    bytes_map = {b['name']: b['file_bytes'] for b in banners}

    for idx, entry in enumerate(effective_matched):
        if progress_callback:
            progress_callback(idx / total, f"Processing {entry['q_id']}…")

        try:
            q_id_key  = entry['q_id']
            xl_key    = q_id_key[:31]

            # Parse each banner's sheet for this question
            merged_col_labels  = []
            merged_base_values = []
            merged_answers     = None   # use first banner's answers as reference
            merged_data_cols   = []     # list of per-answer value lists, one per column

            for b in banners:
                bname      = b['name']
                sheet_meta = entry['banner_sheets'].get(bname)
                sel_cols   = selected_cols_per_banner.get(bname, [])

                if sheet_meta is None or not sel_cols:
                    # Banner doesn't have this question — add blank columns
                    for col in sel_cols:
                        merged_col_labels.append(col)
                        merged_base_values.append(None)
                        # Will fill with None per answer below
                        merged_data_cols.append(None)  # sentinel
                    continue

                # Parse the sheet
                fb  = bytes_map[bname]
                xl  = pd.ExcelFile(io.BytesIO(fb))
                df  = xl.parse(sheet_meta['index'], header=None, na_values=[''])
                fmt = detect_format(df)

                if fmt == 'fmt2':
                    parsed = parse_fmt2_sheet(df, sel_cols, weighted_data)
                elif fmt == 'fmt5':
                    parsed = parse_fmt5_sheet(df, sel_cols, weighted_data)
                else:
                    # fmt3/fmt4 not merged in Scenario A for now
                    skipped.append(f"{bname} — {sheet_meta['sheet_name']} (fmt not supported in merge)")
                    for col in sel_cols:
                        merged_col_labels.append(col)
                        merged_base_values.append(None)
                        merged_data_cols.append(None)
                    continue

                col_names = [g for (g, s) in parsed['columns']]
                answers_b = parsed['answers']
                data_b    = parsed['data']   # list of rows, each row = list of values per col

                # Set reference answers from first banner
                if merged_answers is None:
                    merged_answers = answers_b

                # Build answer → data lookup for this banner
                ans_lookup = {}
                for ai, ans in enumerate(answers_b):
                    ans_lookup[ans.strip().lower()] = data_b[ai] if ai < len(data_b) else []

                for ci, col_name in enumerate(col_names):
                    merged_col_labels.append(col_name)
                    merged_base_values.append(
                        parsed['base_values'][ci] if ci < len(parsed['base_values']) else None
                    )
                    # Build column data aligned to merged_answers
                    col_data = []
                    for ans in (merged_answers or []):
                        row = ans_lookup.get(ans.strip().lower())
                        val = row[ci] if (row and ci < len(row)) else None
                        col_data.append(val)
                    merged_data_cols.append(col_data)

            if not merged_answers:
                skipped.append(q_id_key)
                continue

            # Fill None sentinel columns with blanks aligned to merged_answers
            n_answers = len(merged_answers)
            for ci in range(len(merged_data_cols)):
                if merged_data_cols[ci] is None:
                    merged_data_cols[ci] = [None] * n_answers

            # Transpose: merged_data[answer_idx][col_idx]
            merged_data = []
            for ai in range(n_answers):
                row = [merged_data_cols[ci][ai] if ai < len(merged_data_cols[ci])
                       else None for ci in range(len(merged_col_labels))]
                merged_data.append(row)

            # Apply row filter
            filtered_answers, filtered_data = _apply_row_filter(merged_answers, merged_data)
            if not filtered_answers:
                skipped.append(q_id_key)
                continue

            question_wording = entry['wording']

            # Write to Word
            if word_doc is not None:
                write_table_to_doc(
                    word_doc, question_wording,
                    merged_col_labels, merged_base_values,
                    filtered_answers, filtered_data, 100, first_doc,
                )
                first_doc = False

            # Write to Excel
            if xl_wb is not None:
                ws, row_num = _xl_get_sheet(xl_key)
                next_row = _xl_write_table(
                    ws, row_num, question_wording,
                    merged_col_labels, merged_base_values,
                    filtered_answers, filtered_data, 100, True,
                )
                _xl_done(ws, xl_key, next_row)

        except Exception as e:
            errors.append((entry.get('wording', entry['q_id'])[:60], str(e)))

    if progress_callback:
        progress_callback(1.0, "Finalizing…")

    word_bytes  = None
    excel_bytes = None

    if word_doc is not None and output_format in ('word', 'both'):
        buf = io.BytesIO()
        word_doc.save(buf)
        word_bytes = buf.getvalue()

    if xl_wb is not None and output_format in ('excel', 'both'):
        buf = io.BytesIO()
        xl_wb.save(buf)
        excel_bytes = buf.getvalue()

    return {
        'word_bytes':  word_bytes,
        'excel_bytes': excel_bytes,
        'skipped':     skipped,
        'errors':      errors,
    }
